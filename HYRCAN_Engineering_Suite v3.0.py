# -*- coding: utf-8 -*-
"""
HYRCAN Engineering Suite v4.0
Developer : Mathias Adjei Tawiah
Institution: Hohai University — College of Harbor, Coastal & Offshore Engineering
Standards  : EurOtop 2018 | JTS 154-1-2011 | GB 50286-2013
"""

import streamlit as st
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.lines as mlines
import numpy as np
import json, io, math, os, datetime
import pandas as pd

# ── optional heavy imports (graceful degradation) ────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAS_XL = True
except ImportError:
    HAS_XL = False

# ════════════════════════════════════════════════════════════════════
#  PAGE CONFIG  (must be first Streamlit call)
# ════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="HYRCAN Engineering Suite v4.0",
    page_icon="⚓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ════════════════════════════════════════════════════════════════════
#  GLOBAL STYLE — premium dark engineering dashboard
# ════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap');

html, body, [class*="css"] { font-family: 'Inter', 'Segoe UI', sans-serif; font-size:14px; }
.stApp { background:#07090f; color:#d1d9e6; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    background:#0d1117; border-bottom:1px solid #1c2230; gap:0; padding:0 4px;
}
.stTabs [data-baseweb="tab"] {
    background:transparent; color:#7d8590; border-radius:0;
    font-weight:600; font-size:13px; padding:11px 18px;
    border-bottom:2px solid transparent; letter-spacing:.2px;
    transition: color .15s, border-color .15s;
}
.stTabs [aria-selected="true"] {
    background:transparent !important; color:#e6edf3 !important;
    border-bottom:2px solid #4493f8 !important;
}
/* ── Inputs ── */
.stNumberInput input, .stTextInput input {
    background:#0d1117 !important; border:1px solid #21262d !important;
    color:#e6edf3 !important; border-radius:6px !important;
    font-size:13px !important; font-family:'Inter',sans-serif !important;
}
.stNumberInput input:focus, .stTextInput input:focus {
    border-color:#4493f8 !important;
    box-shadow:0 0 0 3px rgba(68,147,248,.15) !important;
}
/* ── Buttons ── */
.stButton > button {
    background:#1a7f37; color:#fff; border:1px solid #238636;
    border-radius:6px; font-weight:600; font-size:13px;
    padding:8px 18px; transition:all .15s; letter-spacing:.2px;
}
.stButton > button:hover { background:#238636; border-color:#2ea043; }
/* ── Download buttons ── */
.stDownloadButton > button {
    background:#0d1117 !important; color:#4493f8 !important;
    border:1px solid #21262d !important; border-radius:6px !important;
    font-weight:500 !important; font-size:12px !important;
    padding:7px 14px !important; transition:all .15s !important;
}
.stDownloadButton > button:hover {
    background:#161b22 !important; border-color:#4493f8 !important;
}
/* ── Metrics ── */
[data-testid="metric-container"] {
    background:#0d1117; border:1px solid #21262d; border-radius:8px; padding:14px;
}
[data-testid="metric-container"] [data-testid="metric-label"] {
    font-size:11px !important; color:#7d8590 !important;
    font-weight:600 !important; letter-spacing:.8px !important; text-transform:uppercase !important;
}
[data-testid="metric-container"] [data-testid="metric-value"] {
    font-size:22px !important; font-weight:700 !important; color:#e6edf3 !important;
    font-family:'JetBrains Mono',monospace !important;
}
/* ── DataFrames ── */
.stDataFrame { border-radius:8px; overflow:hidden; }
.stDataFrame thead tr th { font-size:11px !important; font-weight:700 !important; letter-spacing:.6px !important; text-transform:uppercase !important; }
.stAlert { border-radius:6px; font-size:13px; }
/* ── Sidebar ── */
[data-testid="stSidebar"] { background:#0b0f18; border-right:1px solid #1c2230; }

/* ── Section labels ── */
.sec-lbl {
    display:flex; align-items:center; gap:8px;
    font-size:10px; font-weight:700; letter-spacing:1.2px; text-transform:uppercase;
    color:#7d8590; margin:22px 0 10px 0; padding-bottom:8px;
    border-bottom:1px solid #1c2230;
}
.sec-lbl .sl-dot {
    width:5px; height:5px; border-radius:50%;
    background:#4493f8; flex-shrink:0;
}

/* ── Code/mono output blocks ── */
.code-out {
    background:#0d1117; border:1px solid #21262d; border-radius:8px;
    padding:18px 20px; font-family:'JetBrains Mono','Consolas',monospace;
    font-size:12px; color:#d1d9e6; line-height:1.75;
    overflow-x:auto; white-space:pre;
}

/* ── Equation blocks ── */
.eq-card {
    background:#0d1117; border:1px solid #21262d;
    border-left:3px solid #4493f8; border-radius:0 8px 8px 0;
    padding:14px 18px; margin:10px 0;
}
.eq-step { font-size:10px; font-weight:700; letter-spacing:1px; text-transform:uppercase; color:#4493f8; margin-bottom:6px; }
.eq-formula { font-family:'JetBrains Mono',monospace; font-size:15px; color:#e6edf3; font-weight:500; margin-bottom:5px; }
.eq-sub { font-family:'JetBrains Mono',monospace; font-size:12px; color:#7d8590; margin-bottom:5px; white-space:pre; }
.eq-result { font-family:'JetBrains Mono',monospace; font-size:14px; color:#3fb950; font-weight:700; margin-top:6px; }
.eq-ref { font-size:10px; color:#484f58; margin-top:4px; }

/* ── Material properties table ── */
.mat-tbl { width:100%; border-collapse:collapse; font-size:13px; }
.mat-tbl th {
    background:#0d1117; color:#7d8590; font-size:10px; font-weight:700;
    letter-spacing:.8px; text-transform:uppercase; padding:10px 14px;
    border-bottom:2px solid #4493f8; text-align:left;
}
.mat-tbl td { padding:10px 14px; border-bottom:1px solid #1c2230; color:#d1d9e6; vertical-align:middle; }
.mat-tbl tr:last-child td { border-bottom:none; }
.mat-tbl tr:nth-child(even) td { background:#0b0f18; }
.mat-tbl tr:hover td { background:#161b22; }
.mat-tbl .mn { font-family:'JetBrains Mono',monospace; text-align:right; }
.mat-tbl .lb { font-weight:600; color:#e6edf3; }

/* ── Status pills ── */
.p-pass { background:#12261a; color:#3fb950; border:1px solid #238636; border-radius:12px; padding:2px 10px; font-size:11px; font-weight:700; }
.p-fail { background:#27100e; color:#f85149; border:1px solid #a82726; border-radius:12px; padding:2px 10px; font-size:11px; font-weight:700; }
.p-warn { background:#251c08; color:#d29922; border:1px solid #9e6a03; border-radius:12px; padding:2px 10px; font-size:11px; font-weight:700; }

/* ── Info / history cards ── */
.info-card {
    background:#0d1117; border:1px solid #21262d; border-radius:8px;
    padding:14px 18px; margin-bottom:10px;
}
.ic-title { font-size:10px; font-weight:700; letter-spacing:1px; text-transform:uppercase; color:#7d8590; margin-bottom:8px; }
.ic-body { font-size:12px; color:#7d8590; line-height:2; }

/* ── Brand header ── */
.brand-strip {
    background:linear-gradient(135deg,#0f1824 0%,#0d1117 60%,#0a1020 100%);
    border-bottom:1px solid #1c2230;
    padding:14px 0 14px 0; margin-bottom:2px;
    display:flex; align-items:center; gap:14px;
}

/* ── History record card ── */
.hist-card {
    background:#0d1117; border:1px solid #21262d; border-radius:8px;
    padding:14px 18px; margin-bottom:8px; transition:border-color .15s;
}
.hist-card:hover { border-color:#4493f8; }

/* ── Sidebar link ── */
.sb-link a { color:#4493f8 !important; text-decoration:none; font-size:12px; line-height:2.2; }
.sb-link a:hover { text-decoration:underline; }

/* ── Badge ── */
.badge {
    display:inline-block; background:#122038; border:1px solid #1e4070;
    color:#4493f8; border-radius:20px; padding:2px 10px;
    font-size:11px; font-weight:600; margin:2px;
}
</style>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════
#  PROJECT HISTORY  (persists within server session via /tmp)
# ════════════════════════════════════════════════════════════════════
HIST_FILE = "/tmp/hyrcan_v4_history.json"

def load_history():
    if os.path.exists(HIST_FILE):
        try:
            with open(HIST_FILE) as f: return json.load(f)
        except Exception: return []
    return []

def save_history(h):
    with open(HIST_FILE, "w") as f: json.dump(h, f, indent=2)

def push_history(module, label, params, results):
    h = load_history()
    h.insert(0, {
        "id": datetime.datetime.now().strftime("%Y%m%d_%H%M%S"),
        "ts": datetime.datetime.now().strftime("%Y-%m-%d  %H:%M"),
        "module": module, "label": label,
        "params": params, "results": results,
    })
    save_history(h[:60])

# ════════════════════════════════════════════════════════════════════
#  SESSION DEFAULTS
# ════════════════════════════════════════════════════════════════════
def _d(k, v):
    if k not in st.session_state: st.session_state[k] = v

_d('rm_crest_elev', 14.55); _d('rm_seabed_elev', -3.30)
_d('rm_crest_width', 6.0);  _d('rm_upper_height', 4.85)
_d('rm_lower_height', 13.0); _d('rm_sw_upper', 2.5)
_d('rm_sw_lower', 2.5);     _d('rm_sw_berm', 6.5)
_d('rm_lw_upper', 1.5);     _d('rm_lw_lower', 1.5)
_d('rm_lw_berm', 2.0);      _d('rm_dhw', 4.99)
_d('rm_surcharge', 10.0);   _d('rm_num_slices', 50.0)
_d('rm_n_layers', 3)

for i,nm,t,g,c,phi in [(1,'Soft Silt',7.8,16.5,6.0,8.0),
                        (2,'Silty Clay',10.2,19.0,20.0,18.0),
                        (3,'Fine Sand',10.0,19.5,0.0,30.0),
                        (4,'Dense Sand',8.0,20.0,0.0,35.0),
                        (5,'Gravel',6.0,21.0,0.0,40.0)]:
    _d(f'rm_n{i}',nm); _d(f'rm_t{i}',t)
    _d(f'rm_g{i}',g);  _d(f'rm_c{i}',c); _d(f'rm_phi{i}',phi)
_d('rm_gr',19.0); _d('rm_cr_',0.0); _d('rm_phir',38.0)

_d('cs_B',16.0); _d('cs_H_c',14.0); _d('cs_gamma_c',24.0)
_d('cs_d',4.59); _d('cs_H1pct',2.71); _d('cs_gamma_w',10.25)
_d('cs_mu',0.60); _d('cs_q_allow',500.0)

_d('wu_H',2.0); _d('wu_T',10.0); _d('wu_slope',2.5)
_d('wu_Rc',1.5); _d('wu_gamma_f',0.55); _d('wu_gamma_b',1.0)
_d('wu_beta',0.0)
_d('rm_generated', False)

# ════════════════════════════════════════════════════════════════════
#  CORE CALCULATIONS
# ════════════════════════════════════════════════════════════════════
def compute_layer_elevations(seabed_elev, thicknesses):
    tops, bots, cur = [], [], seabed_elev
    for t in thicknesses:
        tops.append(cur); bots.append(cur-t); cur -= t
    datum = abs(bots[-1])
    Y = lambda e: e + datum
    r = {'datum': datum, 'Y_seabed': Y(seabed_elev)}
    for i,(top,bot) in enumerate(zip(tops,bots)):
        r[f'top{i+1}']=top; r[f'bot{i+1}']=bot
        r[f'Y_top{i+1}']=Y(top); r[f'Y_bot{i+1}']=Y(bot)
    return r

def compute_geometry(ce, se, cw, uh, lh, sw_u, sw_l, sw_b, lw_u, lw_l, lw_b, layers):
    datum=layers['datum']; Y=lambda e: e+datum
    cy=Y(ce); sy=Y(se); single_sw=(sw_b==0); single_lw=(lw_b==0)
    if single_sw:
        x_sw_toe=x_sw_be=x_sw_bs=0.0; x_cl=(uh+lh)*sw_u
    else:
        x_sw_toe=0.0; x_sw_be=lh*sw_l; x_sw_bs=x_sw_be+sw_b; x_cl=x_sw_bs+uh*sw_u
    x_cr=x_cl+cw
    if single_lw:
        x_lw_bs=x_lw_be=x_cr; x_lw_toe=x_cr+(uh+lh)*lw_u
    else:
        x_lw_bs=x_cr+uh*lw_u; x_lw_be=x_lw_bs+lw_b; x_lw_toe=x_lw_be+lh*lw_l
    by=Y(ce-uh)
    lw_by=by if not single_lw else sy
    pts={'sw_toe':(x_sw_toe,sy),'sw_be':(x_sw_be,by if not single_sw else sy),
         'sw_bs':(x_sw_bs,by if not single_sw else sy),
         'cl':(x_cl,cy),'cr':(x_cr,cy),
         'lw_bs':(x_lw_bs,lw_by),'lw_be':(x_lw_be,lw_by),'lw_toe':(x_lw_toe,sy)}
    return {'pts':pts,'crest_y':cy,'seabed_y':sy,'berm_y':by,
            'model_right':x_lw_toe,'model_bottom':0.0,
            'sw_udx':uh*sw_u,'sw_ldx':lh*sw_l,'lw_udx':uh*lw_u,'lw_ldx':lh*lw_l,
            'x_sw_be':x_sw_be,'x_sw_bs':x_sw_bs,'x_cl':x_cl,'x_cr':x_cr,
            'x_lw_bs':x_lw_bs,'x_lw_be':x_lw_be,'x_lw_toe':x_lw_toe,
            'single_slope_sw':single_sw,'single_slope_lw':single_lw}

def verify_geometry(g, ce, se, cw, uh, lh, sw_u, sw_l, sw_b, lw_u, lw_l, lw_b):
    pts=g['pts']; tol=0.001; checks=[]
    def ck(n,e,got): checks.append((n,e,got,abs(e-got)<=tol))
    ck('Crest width',cw,pts['cr'][0]-pts['cl'][0])
    ck('Total height',ce-se,uh+lh)
    if not g['single_slope_sw']:
        ck('Seaward lower ΔX',lh*sw_l,g['sw_ldx'])
        ck('Seaward berm width',sw_b,g['x_sw_bs']-g['x_sw_be'])
        ck('Seaward upper ΔX',uh*sw_u,g['sw_udx'])
    else:
        ck('Seaward single-slope ΔX',(uh+lh)*sw_u,g['x_cl']-pts['sw_toe'][0])
    if not g['single_slope_lw']:
        ck('Landward upper ΔX',uh*lw_u,g['lw_udx'])
        ck('Landward berm width',lw_b,g['x_lw_be']-g['x_lw_bs'])
        ck('Landward lower ΔX',lh*lw_l,g['lw_ldx'])
    else:
        ck('Landward single-slope ΔX',(uh+lh)*lw_u,g['x_lw_toe']-g['x_cr'])
    return checks

def caisson_fos(B,H_c,gamma_c,d,H1pct,gamma_w,mu,q_allow):
    W=B*H_c*gamma_c; P=0.5*gamma_w*H1pct**2+gamma_w*d*H1pct
    F_res=mu*W; M_res=W*(B/2); arm=d/2+H1pct/3; M_ot=P*arm
    return dict(W=W,P=P,F_res=F_res,M_res=M_res,arm=arm,M_ot=M_ot,
                FOS_s=F_res/P if P>0 else float('inf'),
                FOS_o=M_res/M_ot if M_ot>0 else float('inf'),
                q_max=W/B)

def calc_L0(T): return 9.81*T**2/(2*math.pi)
def calc_xi(ta,H,L): return ta/math.sqrt(H/L) if H>0 and L>0 else 0.0
def calc_Ru2(H,xi,gf,gb,gbeta):
    ru=1.65*gb*gf*gbeta*xi*H
    cap=(4.0-1.5/math.sqrt(max(gf*gbeta*xi,1e-12)))*gb*H
    return min(ru,cap)
def calc_Ru1(H,xi,gf,gb,gbeta): return 1.4*calc_Ru2(H,xi,gf,gb,gbeta)
def calc_Cr(xi): return min(0.1*xi**2,1.0)
def calc_q(H,L0,xi,Rc,gf,gbeta):
    if H<=0 or L0<=0 or xi<=0: return 0.0
    t=(Rc/(gf*gbeta*H))*(1.0/xi)
    return max(math.sqrt(9.81*H**3)/math.sqrt(L0/H)*0.2*math.exp(-2.3*t),0.0)
def calc_gbeta(beta):
    b=abs(beta); return max(1.0-0.0033*b,0.736) if b<=80 else 0.736
def breaker(xi):
    if xi<0.5: return "Spilling"
    if xi<2.0: return "Plunging"
    if xi<3.0: return "Collapsing"
    return "Surging"

def build_boundary(g,pts,L,R,B,sy,cy,by,sw_b,lw_b):
    raw=[(L,sy),(L,B),(R,B),(R,sy)]
    if not g['single_slope_lw']:
        if g['lw_ldx']>0.001: raw.append((pts['lw_be'][0],by))
        if lw_b>0.001: raw.append((pts['lw_bs'][0],by))
    raw.append((pts['cr'][0],cy)); raw.append((pts['cl'][0],cy))
    if not g['single_slope_sw']:
        if g['sw_udx']>0.001: raw.append((pts['sw_bs'][0],by))
        if sw_b>0.001: raw.append((pts['sw_be'][0],by))
    raw.append((pts['sw_toe'][0],sy))
    clean=[]
    for pt in raw:
        if not clean or abs(pt[0]-clean[-1][0])>0.001 or abs(pt[1]-clean[-1][1])>0.001:
            clean.append(pt)
    return clean

# ════════════════════════════════════════════════════════════════════
#  HYRCAN INSTRUCTIONS (Step 7 table in HTML, safe to embed)
# ════════════════════════════════════════════════════════════════════
def mat_table_html(mat_rubble, layer_names, layer_props, n_layers):
    rows = [("Rubble Mound (Embankment)", mat_rubble[0], mat_rubble[1], mat_rubble[2], "Embankment body")]
    for i in range(n_layers):
        rows.append((layer_names[i], layer_props[i][0], layer_props[i][1], layer_props[i][2], f"Layer {i+1}"))
    html = """
<table class="mat-tbl">
<thead><tr>
  <th>#</th><th>Material</th>
  <th style="text-align:right">γ (kN/m³)</th>
  <th style="text-align:right">c (kPa)</th>
  <th style="text-align:right">φ (°)</th>
  <th>Region</th>
</tr></thead><tbody>
"""
    for i,(nm,gv,cv,phiv,reg) in enumerate(rows,1):
        html += f"""<tr>
  <td class="mn">{i}</td>
  <td class="lb">{nm}</td>
  <td class="mn">{gv:.1f}</td>
  <td class="mn">{cv:.1f}</td>
  <td class="mn">{phiv:.1f}</td>
  <td style="color:#7d8590;font-size:12px">{reg}</td>
</tr>"""
    html += "</tbody></table>"
    return html

def gen_instructions(g,layers,dhw,q,ns,ce,se,uh,lh,sw_u,sw_l,sw_b,lw_u,lw_l,lw_b,cw,
                     layer_names,layer_props,n_layers,mat_rubble):
    pts=g['pts']; cy=g['crest_y']; sy=g['seabed_y']; by=g['berm_y']
    L=0.0; R=g['model_right']; datum=layers['datum']; dy=dhw+datum
    slope_note="Single-slope" if g['single_slope_sw'] else "Double-slope with berm"
    clean=build_boundary(g,pts,L,R,0.0,sy,cy,by,sw_b,lw_b)
    blines="\n".join([f"  {x:.3f},{y:.3f}" for x,y in clean])+"\n  c"
    bl=""
    for i in range(n_layers-1):
        k=f'Y_bot{i+1}'
        bl+=f"\n  -- Boundary {i+1}: Bottom of {layer_names[i]} / Top of {layer_names[i+1]}\n  {L:.3f},{layers[k]:.3f}\n  {R:.3f},{layers[k]:.3f}\n  d\n"
    ts=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    return f"""\
{'='*70}
  HYRCAN Engineering Suite v4.0  |  Complete Setup Instructions
  Mathias Adjei Tawiah  |  Hohai University
  Generated: {ts}
  Slope: {slope_note}  |  Layers: {n_layers}
{'='*70}

STEP 1 — NEW PROJECT
  File > New Project
  Name : Embankment_Stability
  Units: SI  (kN, m, degrees)

{'─'*70}
STEP 2 — PROJECT SETTINGS
  Analysis > Project Settings

  Failure Direction : Left to Right
  Surface Type      : Circular
  Search Method     : Slope Search
  Methods           : Bishop Simplified, Spencer, GLE/Morgenstern-Price
  Number of slices  : {int(ns)}

{'─'*70}
STEP 3 — EXTERNAL BOUNDARY
  Geometry > External Boundary

{blines}

{'─'*70}
STEP 4 — MATERIAL BOUNDARIES  ({n_layers-1} internal boundaries)
  Geometry > Material Boundary

  -- Boundary 0: Seabed / Top of {layer_names[0]}
  {L:.3f},{layers['Y_seabed']:.3f}
  {R:.3f},{layers['Y_seabed']:.3f}
  d
{bl}
  Verify {n_layers+1} coloured regions visible.

{'─'*70}
STEP 5 — WATER TABLE
  Geometry > Add Water Table

  {L:.3f},{dy:.3f}
  {R:.3f},{dy:.3f}
  d
  (Y = {dy:.3f}  corresponds to real elevation = {dhw:.2f} m)

{'─'*70}
STEP 6 — SURCHARGE LOAD
  Loading > Distributed Load
  Type      : Vertical (Downward)
  Magnitude : {q:.1f} kN/m2
  Crest X   : {pts['cl'][0]:.3f} to {pts['cr'][0]:.3f}

  {pts['cl'][0]:.3f},{cy:.3f}
  {pts['cr'][0]:.3f},{cy:.3f}
  d

{'─'*70}
STEP 7 — MATERIAL PROPERTIES
  Properties > Define Materials
  [See material table in the application UI above — Step 7]

  NOTE: Assign Rubble to the embankment body (above seabed).
  Assign each soil layer to its corresponding depth region.

{'─'*70}
STEP 8 — COMPUTE
  Analysis > Compute  (allow 15-30 seconds)

{'─'*70}
STEP 9 — RESULTS
  Result tab > Record Bishop and Spencer FOS values
  Minimum recommended FOS: 1.30 (static case)

{'='*70}
DATUM REFERENCE
  Bottom of {layer_names[-1]} = {layers[f'bot{n_layers}']:.2f} m (real) -> Y = 0.000 (HYRCAN)
  Datum shift applied: +{datum:.3f} m

MODEL DIMENSIONS
  Total width  : {R:.3f} m
  Crest X range: {pts['cl'][0]:.3f} to {pts['cr'][0]:.3f}  (width = {cw:.3f} m)
  Seaward toe  : X = {pts['sw_toe'][0]:.3f}
  Landward toe : X = {pts['lw_toe'][0]:.3f}
{'='*70}"""

def gen_coords_txt(g,layers,dhw,q,ce,se,n_layers,sw_b=0,lw_b=0):
    pts=g['pts']; cy=g['crest_y']; sy=g['seabed_y']; by=g['berm_y']
    L=0.0; R=g['model_right']; datum=layers['datum']; dy=dhw+datum
    clean=build_boundary(g,pts,L,R,0.0,sy,cy,by,sw_b,lw_b)
    lines=["HYRCAN COORDINATE EXPORT v4.0",f"Datum shift: +{datum:.3f} m","","EXTERNAL BOUNDARY",
           *[f"{x:.3f},{y:.3f}" for x,y in clean],"c","",
           "MATERIAL BOUNDARY 0 (Seabed):",f"{L:.3f},{layers['Y_seabed']:.3f}",
           f"{R:.3f},{layers['Y_seabed']:.3f}","d"]
    for i in range(n_layers-1):
        lines+=["",f"MATERIAL BOUNDARY {i+1}:",f"{L:.3f},{layers[f'Y_bot{i+1}']:.3f}",
                f"{R:.3f},{layers[f'Y_bot{i+1}']:.3f}","d"]
    lines+=["","WATER TABLE:",f"{L:.3f},{dy:.3f}",f"{R:.3f},{dy:.3f}","d","",
            f"SURCHARGE ({q:.1f} kN/m2):",f"{pts['cl'][0]:.3f},{cy:.3f}",
            f"{pts['cr'][0]:.3f},{cy:.3f}","d"]
    return "\n".join(lines)

# ════════════════════════════════════════════════════════════════════
#  WORD EXPORT
# ════════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex6):
    tc=cell._tc; pr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hex6); shd.set(qn('w:val'),'clear')
    pr.append(shd)

def _hdr_row(tbl, cols, bg='003366'):
    row=tbl.add_row()
    for i,v in enumerate(cols):
        c=row.cells[i]; c.text=str(v)
        _set_cell_bg(c,bg)
        for run in c.paragraphs[0].runs:
            run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255)
            run.font.size=Pt(9)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def _data_row(tbl, cols, alt=False):
    row=tbl.add_row()
    fill='EDF2F7' if alt else 'FFFFFF'
    for i,v in enumerate(cols):
        c=row.cells[i]; c.text=str(v)
        _set_cell_bg(c,fill)
        for run in c.paragraphs[0].runs:
            run.font.size=Pt(9)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    return row

def build_word_rubble(g,layers,dhw,q,checks,ce,se,uh,lh,
                      sw_u,sw_l,sw_b,lw_u,lw_l,lw_b,cw,
                      lnames,lprops,nl,mat_rubble,fig_bytes=None):
    if not HAS_DOCX: return io.BytesIO()
    doc=Document()
    for sec in doc.sections:
        sec.left_margin=Cm(2); sec.right_margin=Cm(2)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run("HYRCAN Engineering Suite v4.0")
    run.font.size=Pt(20); run.font.bold=True; run.font.color.rgb=RGBColor(0,70,150)

    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Module 1 — Rubble Mound Coordinate Generator : Technical Report").font.size=Pt(11)

    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d  %H:%M')}  |  Mathias Adjei Tawiah  |  Hohai University").font.size=Pt(9)
    p3.runs[0].font.italic=True

    p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Standards: EurOtop 2018 | JTS 154-1-2011 | GB 50286-2013").font.size=Pt(9)
    doc.add_paragraph()

    # ── Section 1: Geometry ──────────────────────────────────────────
    h=doc.add_heading("1.  Embankment Geometry Parameters",level=1)
    h.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
    _hdr_row(tbl,["Parameter","Value","Unit"])
    params=[("Crest Elevation",f"{ce:.3f}","m"),("Seabed Elevation",f"{se:.3f}","m"),
            ("Total Height",f"{ce-se:.3f}","m"),("Upper Section Height",f"{uh:.3f}","m"),
            ("Lower Section Height",f"{lh:.3f}","m"),("Crest Width",f"{cw:.3f}","m"),
            ("Design High Water Level",f"{dhw:.3f}","m"),("Surcharge Load",f"{q:.1f}","kN/m²"),
            ("Seaward Slope — Upper",f"1:{sw_u:.2f}","H:V"),
            ("Seaward Slope — Lower",f"1:{sw_l:.2f}","H:V"),
            ("Seaward Berm Width",f"{sw_b:.3f}","m"),
            ("Landward Slope — Upper",f"1:{lw_u:.2f}","H:V"),
            ("Landward Slope — Lower",f"1:{lw_l:.2f}","H:V"),
            ("Landward Berm Width",f"{lw_b:.3f}","m"),
            ("Datum Shift Applied",f"+{layers['datum']:.3f}","m")]
    for i,(p_,v,u) in enumerate(params): _data_row(tbl,[p_,v,u],i%2==1)

    doc.add_paragraph()
    h2=doc.add_heading("2.  Material Properties",level=1)
    h2.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl2=doc.add_table(rows=1,cols=5); tbl2.style='Table Grid'
    _hdr_row(tbl2,["#","Material","γ (kN/m³)","c (kPa)","φ (°)"])
    _data_row(tbl2,["1","Rubble Mound (Embankment)",f"{mat_rubble[0]:.1f}",f"{mat_rubble[1]:.1f}",f"{mat_rubble[2]:.1f}"])
    for i in range(nl):
        _data_row(tbl2,[str(i+2),lnames[i],f"{lprops[i][0]:.1f}",f"{lprops[i][1]:.1f}",f"{lprops[i][2]:.1f}"],i%2==1)

    doc.add_paragraph()
    h3=doc.add_heading("3.  Layer Elevation Schedule (HYRCAN Coordinate System)",level=1)
    h3.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl3=doc.add_table(rows=1,cols=6); tbl3.style='Table Grid'
    _hdr_row(tbl3,["Layer","Top Elev (m)","Bot Elev (m)","Thickness (m)","HYRCAN Y-top","HYRCAN Y-bot"])
    for i in range(nl):
        t_=layers[f'top{i+1}']; b_=layers[f'bot{i+1}']
        _data_row(tbl3,[lnames[i],f"{t_:.3f}",f"{b_:.3f}",f"{t_-b_:.3f}",
                        f"{layers[f'Y_top{i+1}']:.3f}",f"{layers[f'Y_bot{i+1}']:.3f}"],i%2==1)

    doc.add_paragraph()
    h4=doc.add_heading("4.  Geometry Verification",level=1)
    h4.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl4=doc.add_table(rows=1,cols=4); tbl4.style='Table Grid'
    _hdr_row(tbl4,["Check","Expected (m)","Calculated (m)","Status"])
    for i,(nm,exp,got,ok) in enumerate(checks):
        r_=_data_row(tbl4,[nm,f"{exp:.4f}",f"{got:.4f}","PASS" if ok else "FAIL"],i%2==1)
        sc=r_.cells[3]
        for run in sc.paragraphs[0].runs:
            run.font.color.rgb=RGBColor(0,128,0) if ok else RGBColor(192,0,0)
            run.font.bold=True

    doc.add_paragraph()
    h5=doc.add_heading("5.  HYRCAN Step-by-Step Setup",level=1)
    h5.runs[0].font.color.rgb=RGBColor(0,70,150)

    # Step-by-step setup as formatted paragraphs
    pts=g['pts']; cy_=g['crest_y']; sy_=g['seabed_y']; by_=g['berm_y']
    datum=layers['datum']; dy_=dhw+datum
    clean=build_boundary(g,pts,0.0,g['model_right'],0.0,sy_,cy_,by_,sw_b,lw_b)
    steps_data=[
        ("STEP 1 — New Project","File > New Project\nName: Embankment_Stability\nUnits: SI (kN, m, degrees)"),
        ("STEP 2 — Project Settings","Failure Direction: Left to Right\nSurface Type: Circular\nSearch Method: Slope Search\nMethods: Bishop Simplified, Spencer, GLE\nNumber of slices: "+str(int(ns if 'ns' in dir() else 50))),
        ("STEP 3 — External Boundary","Geometry > External Boundary\n"+"\n".join([f"{x:.3f},{y:.3f}" for x,y in clean])+"\nc"),
        ("STEP 4 — Material Boundaries","Geometry > Material Boundary\n[Enter seabed boundary and layer boundaries as listed in coordinates export]"),
        ("STEP 5 — Water Table",f"Geometry > Add Water Table\n0.000,{dy_:.3f}\n{g['model_right']:.3f},{dy_:.3f}\nd\n(Y = {dy_:.3f}  ↔  real elevation = {dhw:.2f} m)"),
        ("STEP 6 — Surcharge Load",f"Loading > Distributed Load\nType: Vertical Downward\nMagnitude: {q:.1f} kN/m²\n{pts['cl'][0]:.3f},{cy_:.3f}\n{pts['cr'][0]:.3f},{cy_:.3f}\nd"),
        ("STEP 7 — Material Properties","Properties > Define Materials\n[Use the material table in Section 2 above]\nAssign Rubble to embankment body; assign each soil to its depth region."),
        ("STEP 8 — Compute","Analysis > Compute  (allow 15–30 seconds)"),
        ("STEP 9 — Results","Result tab > Record Bishop and Spencer FOS values\nMinimum required FOS: 1.30 (static)"),
    ]
    for step_title, step_body in steps_data:
        p_s=doc.add_paragraph()
        run=p_s.add_run(step_title)
        run.font.bold=True; run.font.size=Pt(10); run.font.color.rgb=RGBColor(0,70,150)
        doc.add_paragraph(step_body).runs[0].font.size=Pt(9) if doc.paragraphs[-1].runs else None

    if fig_bytes:
        doc.add_paragraph()
        h6=doc.add_heading("6.  Cross-Section Diagram",level=1)
        h6.runs[0].font.color.rgb=RGBColor(0,70,150)
        doc.add_picture(io.BytesIO(fig_bytes), width=Cm(15))

    p_foot=doc.add_paragraph()
    p_foot.add_run("Generated by HYRCAN Engineering Suite v4.0. Verify all parameters against project drawings before use.").font.italic=True
    p_foot.runs[0].font.size=Pt(8)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_word_caisson(cs_p,r,fig_bytes=None):
    if not HAS_DOCX: return io.BytesIO()
    B=cs_p['B']; H_c=cs_p['H_c']; gc=cs_p['gamma_c']; d=cs_p['d']
    H1=cs_p['H1pct']; gw=cs_p['gamma_w']; mu=cs_p['mu']; qa=cs_p['q_allow']
    s_ok=r['FOS_s']>=1.25; o_ok=r['FOS_o']>=1.50; b_ok=r['q_max']<qa
    doc=Document()
    for sec in doc.sections:
        sec.left_margin=Cm(2); sec.right_margin=Cm(2)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run("HYRCAN Engineering Suite v4.0")
    run.font.size=Pt(20); run.font.bold=True; run.font.color.rgb=RGBColor(0,70,150)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Module 2 — Vertical Caisson Stability Analysis : Technical Report").font.size=Pt(11)
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d  %H:%M')}  |  Mathias Adjei Tawiah  |  Hohai University").font.size=Pt(9)
    p3.runs[0].font.italic=True
    doc.add_paragraph()

    h=doc.add_heading("1.  Input Parameters",level=1); h.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
    _hdr_row(tbl,["Parameter","Value","Unit"])
    for i,(p_,v,u) in enumerate([
        ("Caisson Width B",f"{B:.2f}","m"),("Caisson Height H_c",f"{H_c:.2f}","m"),
        ("Concrete Unit Weight γ_c",f"{gc:.2f}","kN/m³"),("Water Depth d",f"{d:.3f}","m"),
        ("Design Wave Height H₁%",f"{H1:.3f}","m"),("Water Unit Weight γ_w",f"{gw:.2f}","kN/m³"),
        ("Friction Coefficient μ",f"{mu:.3f}","—"),
        ("Allowable Bearing Pressure",f"{qa:.0f}","kPa")]): _data_row(tbl,[p_,v,u],i%2==1)

    doc.add_paragraph()
    h2=doc.add_heading("2.  Stability Calculation",level=1); h2.runs[0].font.color.rgb=RGBColor(0,70,150)

    calc_items=[
        ("1.  Self-Weight",
         f"W = B × H_c × γ_c\n"
         f"W = {B:.2f} × {H_c:.2f} × {gc:.2f}\n"
         f"W = {r['W']:,.2f}  kN/m"),
        ("2.  Wave Force  [JTS 154-1-2011, Appendix A]",
         f"P = ½ · γ_w · H₁%² + γ_w · d · H₁%\n"
         f"P = 0.5 × {gw:.2f} × {H1:.2f}² + {gw:.2f} × {d:.2f} × {H1:.2f}\n"
         f"P = {0.5*gw*H1**2:.3f} + {gw*d*H1:.3f} = {r['P']:,.3f}  kN/m"),
        ("3.  Sliding Stability  [GB 50286-2013 §5.3.2]",
         f"FOS_sliding = μ · W / P\n"
         f"FOS_sliding = {mu:.3f} × {r['W']:,.2f} / {r['P']:,.3f}\n"
         f"FOS_sliding = {r['FOS_s']:.4f}    Required ≥ 1.25    → {'PASS' if s_ok else 'FAIL'}"),
        ("4.  Overturning Stability  [GB 50286-2013 §5.3.3]",
         f"arm = d/2 + H₁%/3 = {d/2:.3f} + {H1/3:.3f} = {r['arm']:.4f} m\n"
         f"M_res = W × (B/2) = {r['W']:,.2f} × {B/2:.3f} = {r['M_res']:,.2f}  kN·m/m\n"
         f"M_ot  = P × arm  = {r['P']:,.3f} × {r['arm']:.4f} = {r['M_ot']:,.2f}  kN·m/m\n"
         f"FOS_overturning = {r['M_res']:,.2f} / {r['M_ot']:,.2f} = {r['FOS_o']:.4f}    Required ≥ 1.50    → {'PASS' if o_ok else 'FAIL'}"),
        ("5.  Bearing Pressure  [JTS 154-1-2011 §5.3.4]",
         f"q_max = W / B = {r['W']:,.2f} / {B:.2f} = {r['q_max']:,.2f}  kPa\n"
         f"Allowable = {qa:.0f} kPa    → {'PASS' if b_ok else 'FAIL'}"),
    ]
    for title, body in calc_items:
        p_s=doc.add_paragraph()
        p_s.add_run(title).font.bold=True
        p_s.runs[0].font.size=Pt(10); p_s.runs[0].font.color.rgb=RGBColor(0,70,150)
        pb=doc.add_paragraph(body); pb.runs[0].font.size=Pt(9) if pb.runs else None

    doc.add_paragraph()
    h3=doc.add_heading("3.  Results Summary",level=1); h3.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl2=doc.add_table(rows=1,cols=5); tbl2.style='Table Grid'
    _hdr_row(tbl2,["Check","Calculated","Required","Standard","Status"])
    rows_=[("Self-Weight W",f"{r['W']:,.1f} kN/m","—","—","—"),
           ("Wave Force P",f"{r['P']:,.3f} kN/m","—","JTS 154-1-2011 App. A","—"),
           ("Sliding FOS",f"{r['FOS_s']:.3f}","≥ 1.25","GB 50286-2013 §5.3.2","PASS" if s_ok else "FAIL"),
           ("Overturning FOS",f"{r['FOS_o']:.3f}","≥ 1.50","GB 50286-2013 §5.3.3","PASS" if o_ok else "FAIL"),
           ("Bearing q_max",f"{r['q_max']:,.1f} kPa",f"< {qa:.0f} kPa","JTS 154-1-2011 §5.3.4","PASS" if b_ok else "FAIL")]
    for i,row_d in enumerate(rows_):
        r_=_data_row(tbl2,list(row_d),i%2==1)
        if row_d[4] in ("PASS","FAIL"):
            sc=r_.cells[4]
            for run in sc.paragraphs[0].runs:
                run.font.color.rgb=RGBColor(0,128,0) if row_d[4]=="PASS" else RGBColor(192,0,0)
                run.font.bold=True

    overall="DESIGN ADEQUATE — All stability criteria satisfied." if (s_ok and o_ok and b_ok) \
            else "DESIGN DEFICIENT — One or more criteria not satisfied. Revise parameters."
    doc.add_paragraph()
    p_ov=doc.add_paragraph(overall)
    p_ov.runs[0].font.bold=True; p_ov.runs[0].font.size=Pt(10)
    p_ov.runs[0].font.color.rgb=RGBColor(0,128,0) if (s_ok and o_ok and b_ok) else RGBColor(192,0,0)

    if fig_bytes:
        doc.add_paragraph()
        h4=doc.add_heading("4.  Cross-Section Diagram",level=1); h4.runs[0].font.color.rgb=RGBColor(0,70,150)
        doc.add_picture(io.BytesIO(fig_bytes),width=Cm(14))

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def build_word_runup(wu_in, wu_res):
    if not HAS_DOCX: return io.BytesIO()
    H=wu_in['H']; T=wu_in['T']; slope=wu_in['slope']; Rc=wu_in['Rc']
    gf=wu_in['gf']; gb=wu_in['gb']; beta=wu_in['beta']
    L0=wu_res['L0']; xi=wu_res['xi']; gbeta=wu_res['gbeta']
    Ru2=wu_res['Ru2']; Ru1=wu_res['Ru1']; Cr=wu_res['Cr']; q_wu=wu_res['q']
    btype=wu_res['btype']; ta=1.0/slope

    doc=Document()
    for sec in doc.sections:
        sec.left_margin=Cm(2); sec.right_margin=Cm(2)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run("HYRCAN Engineering Suite v4.0")
    run.font.size=Pt(20); run.font.bold=True; run.font.color.rgb=RGBColor(0,70,150)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Module 3 — Wave Run-Up & Overtopping Calculator : Technical Report").font.size=Pt(11)
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d  %H:%M')}  |  Mathias Adjei Tawiah  |  Hohai University").font.size=Pt(9)
    p3.runs[0].font.italic=True
    doc.add_paragraph()

    h=doc.add_heading("1.  Input Parameters",level=1); h.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
    _hdr_row(tbl,["Parameter","Value","Unit"])
    for i,(p_,v,u) in enumerate([
        ("Significant Wave Height H_m0",f"{H:.3f}","m"),
        ("Peak Wave Period T_p",f"{T:.2f}","s"),
        ("Slope Ratio H:V",f"1:{slope:.2f}","—"),("tan(α)",f"{ta:.4f}","—"),
        ("Freeboard R_c",f"{Rc:.3f}","m"),("Roughness Factor γ_f",f"{gf:.3f}","—"),
        ("Berm Factor γ_b",f"{gb:.3f}","—"),("Wave Obliquity β",f"{beta:.1f}","°")]
    ): _data_row(tbl,[p_,v,u],i%2==1)

    doc.add_paragraph()
    h2=doc.add_heading("2.  Step-by-Step Calculations (EurOtop 2018)",level=1)
    h2.runs[0].font.color.rgb=RGBColor(0,70,150)

    Ru_raw=1.65*gb*gf*gbeta*xi*H
    cap_val=(4.0-1.5/max(math.sqrt(gf*gbeta*xi),1e-9))*gb*H
    calc_items=[
        ("Step 1 — Deep-Water Wavelength",
         f"L₀ = g · T_p² / (2π)\n"
         f"L₀ = 9.81 × {T:.2f}² / (2π)\n"
         f"L₀ = {L0:.4f}  m"),
        ("Step 2 — Iribarren Number (Surf-Similarity Parameter)",
         f"ξ₀ = tan(α) / √(H_m0 / L₀)\n"
         f"ξ₀ = {ta:.4f} / √({H:.3f} / {L0:.3f})\n"
         f"ξ₀ = {xi:.4f}    Breaker classification: {btype}"),
        ("Step 3 — Obliquity Reduction Factor",
         f"γ_β = max(1 − 0.0033 · |β|,  0.736)   for β ≤ 80°\n"
         f"γ_β = max(1 − 0.0033 × {beta:.1f},  0.736)\n"
         f"γ_β = {gbeta:.4f}    [EurOtop 2018, §5.1.3]"),
        ("Step 4 — Wave Run-Up R_u2%  [EurOtop 2018, Eq. 5.2]",
         f"R_u2% = min(1.65 · γ_b · γ_f · γ_β · ξ₀ · H_m0,  (4.0 − 1.5/√(γ_f·γ_β·ξ₀)) · γ_b · H_m0)\n"
         f"Formula value = 1.65 × {gb:.3f} × {gf:.3f} × {gbeta:.4f} × {xi:.4f} × {H:.3f} = {Ru_raw:.4f} m\n"
         f"Cap value     = (4.0 − 1.5/√({gf:.3f}×{gbeta:.4f}×{xi:.4f})) × {gb:.3f} × {H:.3f} = {cap_val:.4f} m\n"
         f"R_u2% = {Ru2:.4f} m     (governing)\n"
         f"R_u1% = 1.4 × R_u2% = {Ru1:.4f} m"),
        ("Step 5 — Reflection Coefficient  [Postma 1989]",
         f"C_r = 0.1 · ξ₀²   (capped at 1.0)\n"
         f"C_r = 0.1 × {xi:.4f}² = {Cr:.4f}"),
        ("Step 6 — Mean Overtopping Discharge  [EurOtop 2018, Eq. 5.6]",
         f"q = √(g·H³) / √(L₀/H) · 0.2 · exp(−2.3 · R_c / (γ_f·γ_β·H·ξ₀))\n"
         f"q = {q_wu:.6f}  m³/s/m\n"
         f"q = {q_wu*1000:.5f}  L/s/m"),
    ]
    for title,body in calc_items:
        p_s=doc.add_paragraph()
        p_s.add_run(title).font.bold=True
        p_s.runs[0].font.size=Pt(10); p_s.runs[0].font.color.rgb=RGBColor(0,70,150)
        pb=doc.add_paragraph(body)
        if pb.runs: pb.runs[0].font.size=Pt(9)

    doc.add_paragraph()
    h3=doc.add_heading("3.  Results Summary",level=1); h3.runs[0].font.color.rgb=RGBColor(0,70,150)
    tbl2=doc.add_table(rows=1,cols=3); tbl2.style='Table Grid'
    _hdr_row(tbl2,["Result Parameter","Value","Unit"])
    for i,(p_,v,u) in enumerate([
        ("Deep-Water Wavelength L₀",f"{L0:.4f}","m"),
        ("Iribarren Number ξ₀",f"{xi:.4f}","—"),
        ("Breaker Classification",btype,"—"),
        ("Obliquity Factor γ_β",f"{gbeta:.4f}","—"),
        ("Run-Up R_u2%",f"{Ru2:.4f}","m"),
        ("Run-Up R_u1%",f"{Ru1:.4f}","m"),
        ("Reflection Coefficient C_r",f"{Cr:.4f}","—"),
        ("Mean Overtopping q",f"{q_wu*1000:.5f}","L/s/m"),
        ("Mean Overtopping q",f"{q_wu:.7f}","m³/s/m"),
        ("R_c / R_u2%",f"{Rc/Ru2:.3f}" if Ru2>0 else "—","—")]
    ): _data_row(tbl2,[p_,v,u],i%2==1)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

# ════════════════════════════════════════════════════════════════════
#  EXCEL EXPORT
# ════════════════════════════════════════════════════════════════════
def _xl_styles():
    thin=Side(style='thin',color='D0D7DE')
    brd=Border(left=thin,right=thin,top=thin,bottom=thin)
    hdr_f=Font(name='Calibri',bold=True,color='FFFFFF',size=10)
    hdr_fill=PatternFill('solid',fgColor='003366')
    dat_f=Font(name='Calibri',size=10)
    alt_fill=PatternFill('solid',fgColor='EEF2F7')
    ctr=Alignment(horizontal='center',vertical='center',wrap_text=True)
    lft=Alignment(horizontal='left',vertical='center')
    return brd,hdr_f,hdr_fill,dat_f,alt_fill,ctr,lft

def _xl_write_table(ws,headers,rows,sr=1,sc=1):
    brd,hf,hfill,df,afill,ctr,lft=_xl_styles()
    for j,h in enumerate(headers,sc):
        c=ws.cell(row=sr,column=j,value=h)
        c.font=hf; c.fill=hfill; c.alignment=ctr; c.border=brd
    for i,row in enumerate(rows,sr+1):
        fill=afill if i%2==0 else PatternFill('solid',fgColor='FFFFFF')
        for j,v in enumerate(row,sc):
            c=ws.cell(row=i,column=j,value=v)
            c.font=df; c.fill=fill; c.alignment=lft; c.border=brd
    return sr+len(rows)+2

def build_excel_rubble(g,layers,dhw,q,checks,ce,se,uh,lh,
                       sw_u,sw_l,sw_b,lw_u,lw_l,lw_b,cw,
                       lnames,lprops,nl,mat_rubble):
    if not HAS_XL: return io.BytesIO()
    wb=openpyxl.Workbook(); ts=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    title_f=Font(name='Calibri',bold=True,size=14,color='003366')
    sub_f=Font(name='Calibri',bold=True,size=11,color='004499')
    meta_f=Font(name='Calibri',italic=True,size=9,color='7D8590')
    pass_f=Font(name='Calibri',bold=True,size=10,color='006100')
    fail_f=Font(name='Calibri',bold=True,size=10,color='9C0006')
    pass_fill=PatternFill('solid',fgColor='C6EFCE')
    fail_fill=PatternFill('solid',fgColor='FFC7CE')

    # Sheet 1: Geometry
    ws=wb.active; ws.title="Geometry"
    ws.sheet_view.showGridLines=False
    ws['A1']="HYRCAN v4.0 — Rubble Mound Module"; ws['A1'].font=title_f; ws.merge_cells('A1:E1')
    ws['A2']=f"Generated: {ts}  |  Mathias Adjei Tawiah  |  Hohai University"; ws['A2'].font=meta_f; ws.merge_cells('A2:E2')
    ws['A4']="EMBANKMENT GEOMETRY"; ws['A4'].font=sub_f
    _xl_write_table(ws,["Parameter","Value","Unit"],[
        ["Crest Elevation",ce,"m"],["Seabed Elevation",se,"m"],
        ["Total Height",round(ce-se,3),"m"],["Upper Section Height",uh,"m"],
        ["Lower Section Height",lh,"m"],["Crest Width",cw,"m"],
        ["Design High Water Level",dhw,"m"],["Surcharge Load",q,"kN/m²"],
        ["Seaward Slope — Upper",f"1:{sw_u:.2f}","H:V"],["Seaward Slope — Lower",f"1:{sw_l:.2f}","H:V"],
        ["Seaward Berm Width",sw_b,"m"],["Landward Slope — Upper",f"1:{lw_u:.2f}","H:V"],
        ["Landward Slope — Lower",f"1:{lw_l:.2f}","H:V"],["Landward Berm Width",lw_b,"m"],
        ["Datum Shift",round(layers['datum'],3),"m"]
    ],sr=5)
    ws.column_dimensions['A'].width=34; ws.column_dimensions['B'].width=18; ws.column_dimensions['C'].width=14

    # Sheet 2: Layer Schedule
    ws2=wb.create_sheet("Layer Schedule"); ws2.sheet_view.showGridLines=False
    ws2['A1']="HYRCAN v4.0 — Layer Elevation Schedule"; ws2['A1'].font=title_f; ws2.merge_cells('A1:H1')
    ws2['A2']=f"Generated: {ts}"; ws2['A2'].font=meta_f; ws2.merge_cells('A2:H2')
    rows_lyr=[]
    for i in range(nl):
        t_=layers[f'top{i+1}']; b_=layers[f'bot{i+1}']
        rows_lyr.append([lnames[i],lprops[i][0],lprops[i][1],lprops[i][2],
                         round(t_,3),round(b_,3),round(t_-b_,3),
                         round(layers[f'Y_top{i+1}'],3),round(layers[f'Y_bot{i+1}'],3)])
    _xl_write_table(ws2,["Layer Name","γ (kN/m³)","c (kPa)","φ (°)",
                         "Top Elev (m)","Bot Elev (m)","Thickness (m)",
                         "HYRCAN Y-top","HYRCAN Y-bot"],rows_lyr,sr=4)
    for col in ['A','B','C','D','E','F','G','H','I']:
        ws2.column_dimensions[col].width=17

    # Sheet 3: Materials
    ws3=wb.create_sheet("Materials"); ws3.sheet_view.showGridLines=False
    ws3['A1']="HYRCAN v4.0 — Material Properties"; ws3['A1'].font=title_f; ws3.merge_cells('A1:E1')
    ws3['A2']=f"Generated: {ts}"; ws3['A2'].font=meta_f; ws3.merge_cells('A2:E2')
    mat_rows=[["Rubble Mound (Embankment)",mat_rubble[0],mat_rubble[1],mat_rubble[2],"Embankment body"]]
    for i in range(nl): mat_rows.append([lnames[i],lprops[i][0],lprops[i][1],lprops[i][2],f"Layer {i+1}"])
    _xl_write_table(ws3,["Material","γ (kN/m³)","c (kPa)","φ (°)","Region Assignment"],mat_rows,sr=4)
    ws3.column_dimensions['A'].width=30; ws3.column_dimensions['E'].width=22

    # Sheet 4: Verification
    ws4=wb.create_sheet("Verification"); ws4.sheet_view.showGridLines=False
    ws4['A1']="HYRCAN v4.0 — Geometry Verification"; ws4['A1'].font=title_f; ws4.merge_cells('A1:E1')
    ws4['A2']=f"Generated: {ts}"; ws4['A2'].font=meta_f; ws4.merge_cells('A2:E2')
    end_r=_xl_write_table(ws4,["Check","Expected (m)","Calculated (m)","Δ","Status"],
        [[nm,round(exp,4),round(got,4),round(abs(exp-got),5),"PASS" if ok else "FAIL"]
         for nm,exp,got,ok in checks],sr=4)
    for i,(nm,exp,got,ok) in enumerate(checks,5):
        c=ws4.cell(row=i,column=5)
        c.font=pass_f if ok else fail_f
        c.fill=pass_fill if ok else fail_fill
    ws4.column_dimensions['A'].width=32
    for col in ['B','C','D','E']: ws4.column_dimensions[col].width=18

    # Sheet 5: Coordinates
    ws5=wb.create_sheet("Coordinates"); ws5.sheet_view.showGridLines=False
    ws5['A1']="HYRCAN v4.0 — Boundary Coordinates"; ws5['A1'].font=title_f; ws5.merge_cells('A1:B1')
    pts=g['pts']; cy_=g['crest_y']; sy_=g['seabed_y']; by_=g['berm_y']
    clean=build_boundary(g,pts,0.0,g['model_right'],0.0,sy_,cy_,by_,0,0)
    _xl_write_table(ws5,["X (m)","Y (m)"],[[round(x,4),round(y,4)] for x,y in clean],sr=3)
    ws5.column_dimensions['A'].width=20; ws5.column_dimensions['B'].width=20

    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf

def build_excel_caisson(cs_p,r):
    if not HAS_XL: return io.BytesIO()
    B=cs_p['B']; H_c=cs_p['H_c']; gc=cs_p['gamma_c']; d=cs_p['d']
    H1=cs_p['H1pct']; gw=cs_p['gamma_w']; mu=cs_p['mu']; qa=cs_p['q_allow']
    s_ok=r['FOS_s']>=1.25; o_ok=r['FOS_o']>=1.50; b_ok=r['q_max']<qa
    wb=openpyxl.Workbook(); ts=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    title_f=Font(name='Calibri',bold=True,size=14,color='003366')
    meta_f=Font(name='Calibri',italic=True,size=9,color='7D8590')
    pass_f=Font(name='Calibri',bold=True,size=10,color='006100')
    fail_f=Font(name='Calibri',bold=True,size=10,color='9C0006')
    pass_fill=PatternFill('solid',fgColor='C6EFCE')
    fail_fill=PatternFill('solid',fgColor='FFC7CE')

    ws=wb.active; ws.title="Caisson Stability"; ws.sheet_view.showGridLines=False
    ws['A1']="HYRCAN v4.0 — Vertical Caisson Stability"; ws['A1'].font=title_f; ws.merge_cells('A1:E1')
    ws['A2']=f"Generated: {ts}  |  Mathias Adjei Tawiah  |  Hohai University"; ws['A2'].font=meta_f; ws.merge_cells('A2:E2')
    ws['A4']="INPUT PARAMETERS"; ws['A4'].font=Font(name='Calibri',bold=True,size=11,color='004499')
    _xl_write_table(ws,["Parameter","Value","Unit"],[
        ["Caisson Width B",B,"m"],["Caisson Height H_c",H_c,"m"],
        ["Concrete Unit Weight γ_c",gc,"kN/m³"],["Water Depth d",d,"m"],
        ["Design Wave Height H₁%",H1,"m"],["Water Unit Weight γ_w",gw,"kN/m³"],
        ["Friction Coefficient μ",mu,"—"],["Allowable Bearing q_allow",qa,"kPa"]],sr=5)
    ws['A15']="STABILITY RESULTS"; ws['A15'].font=Font(name='Calibri',bold=True,size=11,color='004499')
    end=_xl_write_table(ws,["Check","Calculated","Required","Standard","Status"],[
        ["Self-Weight W",f"{r['W']:,.1f} kN/m","—","—","—"],
        ["Wave Force P",f"{r['P']:,.3f} kN/m","—","JTS 154-1-2011 App. A","—"],
        ["Sliding FOS",round(r['FOS_s'],4),"≥ 1.25","GB 50286-2013 §5.3.2","PASS" if s_ok else "FAIL"],
        ["Overturning FOS",round(r['FOS_o'],4),"≥ 1.50","GB 50286-2013 §5.3.3","PASS" if o_ok else "FAIL"],
        ["Bearing q_max",f"{r['q_max']:,.1f} kPa",f"< {qa:.0f} kPa","JTS 154-1-2011 §5.3.4","PASS" if b_ok else "FAIL"]],sr=16)
    for i,(nm,calc,req,std,st_) in enumerate([
        ("","","","",""),("","","","",""),
        ("","","","","PASS" if s_ok else "FAIL"),
        ("","","","","PASS" if o_ok else "FAIL"),
        ("","","","","PASS" if b_ok else "FAIL")],17):
        if st_ in ("PASS","FAIL"):
            c=ws.cell(row=i,column=5)
            c.font=pass_f if st_=="PASS" else fail_f
            c.fill=pass_fill if st_=="PASS" else fail_fill
    ws.column_dimensions['A'].width=28; ws.column_dimensions['B'].width=22
    ws.column_dimensions['C'].width=16; ws.column_dimensions['D'].width=30; ws.column_dimensions['E'].width=14
    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf

def build_excel_runup(wu_in,wu_res):
    if not HAS_XL: return io.BytesIO()
    wb=openpyxl.Workbook(); ts=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    title_f=Font(name='Calibri',bold=True,size=14,color='003366')
    meta_f=Font(name='Calibri',italic=True,size=9,color='7D8590')
    sub_f=Font(name='Calibri',bold=True,size=11,color='004499')
    ws=wb.active; ws.title="Wave Run-Up"; ws.sheet_view.showGridLines=False
    ws['A1']="HYRCAN v4.0 — Wave Run-Up & Overtopping"; ws['A1'].font=title_f; ws.merge_cells('A1:C1')
    ws['A2']=f"Generated: {ts}  |  Mathias Adjei Tawiah  |  Hohai University"; ws['A2'].font=meta_f; ws.merge_cells('A2:C2')
    ws['A3']="Standard: EurOtop 2018 (Van der Meer et al.)"; ws['A3'].font=meta_f; ws.merge_cells('A3:C3')
    ws['A5']="INPUT PARAMETERS"; ws['A5'].font=sub_f
    _xl_write_table(ws,["Parameter","Value","Unit"],[
        ["Significant Wave Height H_m0",wu_in['H'],"m"],
        ["Peak Wave Period T_p",wu_in['T'],"s"],
        ["Slope Ratio H:V",f"1:{wu_in['slope']:.2f}","—"],
        ["tan(α)",round(1/wu_in['slope'],5),"—"],
        ["Freeboard R_c",wu_in['Rc'],"m"],
        ["Roughness Factor γ_f",wu_in['gf'],"—"],
        ["Berm Factor γ_b",wu_in['gb'],"—"],
        ["Wave Obliquity β",wu_in['beta'],"°"]],sr=6)
    ws['A16']="COMPUTED RESULTS"; ws['A16'].font=sub_f
    _xl_write_table(ws,["Result Parameter","Value","Unit"],[
        ["Deep-Water Wavelength L₀",round(wu_res['L0'],4),"m"],
        ["Iribarren Number ξ₀",round(wu_res['xi'],5),"—"],
        ["Breaker Classification",wu_res['btype'],"—"],
        ["Obliquity Factor γ_β",round(wu_res['gbeta'],5),"—"],
        ["Run-Up R_u2%",round(wu_res['Ru2'],5),"m"],
        ["Run-Up R_u1%",round(wu_res['Ru1'],5),"m"],
        ["Reflection Coefficient C_r",round(wu_res['Cr'],5),"—"],
        ["Mean Overtopping q",round(wu_res['q'],8),"m³/s/m"],
        ["Mean Overtopping q",round(wu_res['q']*1000,6),"L/s/m"],
        ["R_c / R_u2%",round(wu_res['Rc']/wu_res['Ru2'],4) if wu_res['Ru2']>0 else "—","—"]],sr=17)
    ws.column_dimensions['A'].width=38; ws.column_dimensions['B'].width=22; ws.column_dimensions['C'].width=14
    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf

# ════════════════════════════════════════════════════════════════════
#  VISUALIZATIONS
# ════════════════════════════════════════════════════════════════════
LAYER_COLORS=['#C9B99A','#A8B5A0','#B5C4D1','#D1B5C4','#B5D1C4']
BG='#07090f'

def draw_rubble(g,layers,dhw,ce,se,uh,lh,lnames,nl,q_val=10.0):
    pts=g['pts']; berm_e=ce-uh
    fig,ax=plt.subplots(figsize=(12,6),dpi=120)
    fig.patch.set_facecolor(BG); ax.set_facecolor('#0b0f18')
    xl=pts['sw_toe'][0]; xr=pts['lw_toe'][0]
    for i in range(nl):
        t_=layers[f'top{i+1}']; b_=layers[f'bot{i+1}']
        ax.fill_betweenx([b_,t_],xl-999,xr+999,
                          facecolor=LAYER_COLORS[i%len(LAYER_COLORS)],alpha=.50-i*.02,zorder=1)
        ax.axhline(y=t_,color='#2a3040',lw=.8,ls=':',zorder=2)
    ax.axhline(y=se,color='#607d8b',lw=1.3,ls='--',zorder=2)
    if g['single_slope_sw'] and g['single_slope_lw']:
        xs=[pts['sw_toe'][0],pts['cl'][0],pts['cr'][0],pts['lw_toe'][0],pts['sw_toe'][0]]
        ys=[se,ce,ce,se,se]
    elif g['single_slope_sw']:
        xs=[pts['sw_toe'][0],pts['cl'][0],pts['cr'][0],pts['lw_bs'][0],pts['lw_be'][0],pts['lw_toe'][0],pts['sw_toe'][0]]
        ys=[se,ce,ce,berm_e,berm_e,se,se]
    elif g['single_slope_lw']:
        xs=[pts['sw_toe'][0],pts['sw_be'][0],pts['sw_bs'][0],pts['cl'][0],pts['cr'][0],pts['lw_toe'][0],pts['sw_toe'][0]]
        ys=[se,berm_e,berm_e,ce,ce,se,se]
    else:
        xs=[pts['sw_toe'][0],pts['sw_be'][0],pts['sw_bs'][0],pts['cl'][0],pts['cr'][0],
            pts['lw_bs'][0],pts['lw_be'][0],pts['lw_toe'][0],pts['sw_toe'][0]]
        ys=[se,berm_e,berm_e,ce,ce,berm_e,berm_e,se,se]
    ax.fill(xs,ys,facecolor='#C9A96E',alpha=.88,edgecolor='#8B6A38',lw=1.8,zorder=3)
    ax.axhline(y=dhw,color='#4493f8',ls='--',lw=2.0,zorder=4)
    cl_x=pts['cl'][0]; cr_x=pts['cr'][0]; cw_=cr_x-cl_x
    bh=max(cw_*.055,.4)
    ax.add_patch(mpatches.FancyBboxPatch((cl_x,ce+.12),cw_,bh,boxstyle='square,pad=0',
        facecolor='#f7931e22',edgecolor='#f7931e',lw=1.5,hatch='////',zorder=10))
    ax.text((cl_x+cr_x)/2,ce+.12+bh+.55,f'{q_val:.1f} kN/m²',
            ha='center',va='bottom',fontsize=9,fontweight='600',color='#f7931e',zorder=12)
    lx=xl-(xr-xl)*.015
    for i in range(nl):
        ax.text(lx,(layers[f'top{i+1}']+layers[f'bot{i+1}'])/2,lnames[i],
                fontsize=8.5,va='center',ha='right',fontstyle='italic',color='#7d8590',zorder=5)
    handles=[mpatches.Patch(facecolor='#C9A96E',alpha=.88,edgecolor='#8B6A38',label='Rubble Mound'),
             *[mpatches.Patch(facecolor=LAYER_COLORS[i%len(LAYER_COLORS)],alpha=.7,
                label=f'{lnames[i]}  [{layers[f"top{i+1}"]:.1f}→{layers[f"bot{i+1}"]:.1f} m]')
               for i in range(nl)],
             mlines.Line2D([],[],color='#4493f8',ls='--',lw=2,label=f'DHW ({dhw:.2f} m)')]
    leg=ax.legend(handles=handles,loc='upper center',bbox_to_anchor=(.5,-.13),ncol=3,
                  fontsize=9,frameon=True,facecolor='#0d1117',edgecolor='#1c2230',framealpha=.97)
    for t in leg.get_texts(): t.set_color('#d1d9e6')
    ax.set_xlabel('Distance  (m)',fontsize=11,color='#7d8590')
    ax.set_ylabel('Elevation  (m)',fontsize=11,color='#7d8590')
    ax.set_title('Rubble Mound Embankment — Cross-Section',fontsize=13,fontweight='bold',color='#e6edf3',pad=12)
    ax.grid(True,alpha=.10,color='#2a3040',lw=.5)
    ax.tick_params(colors='#7d8590',labelsize=9)
    for sp in ax.spines.values(): sp.set_color('#1c2230')
    mx=(xr-xl)*.08; ax.set_xlim(xl-mx*3.5,xr+mx)
    ax.set_ylim(layers[f'bot{nl}']-1.5,ce+4.5); ax.set_aspect('equal',adjustable='box')
    fig.subplots_adjust(bottom=.22)
    return fig

def draw_caisson(p,r):
    fig,ax=plt.subplots(figsize=(10,7),dpi=120)
    fig.patch.set_facecolor(BG); ax.set_facecolor('#0a0f18')
    B=p['B']; H_c=p['H_c']; d=p['d']
    rb_h,rb_ext=1.8,3.0
    ax.fill([-rb_ext,-rb_ext*.3,B+rb_ext*.3,B+rb_ext,-rb_ext],
             [-rb_h,0,0,-rb_h,-rb_h],
             facecolor='#C8A96E',alpha=.85,edgecolor='#8B6A38',lw=1.2,zorder=2,label='Rubble Foundation')
    ax.axhline(y=0,color='#607d8b',lw=1.3,zorder=3)
    ax.add_patch(mpatches.FancyBboxPatch((0,0),B,H_c,boxstyle='square,pad=0',
        facecolor='#1e3a5f',alpha=.9,edgecolor='#2c5f9e',lw=2.0,zorder=4))
    for xi in np.linspace(B*.2,B*.8,4):
        ax.plot([xi,xi],[0,H_c],color='#2a4f80',lw=.5,alpha=.3,zorder=5)
    for yi in [H_c*.25,H_c*.5,H_c*.75]:
        ax.plot([0,B],[yi,yi],color='#2a4f80',lw=.5,alpha=.3,zorder=5)
    ax.text(B/2,H_c/2,'CAISSON\n(Concrete)',ha='center',va='center',
            fontsize=11,fontweight='bold',color='#d1d9e6',zorder=6,alpha=.9)
    ax.axhline(y=d,color='#4493f8',ls='--',lw=2.0,zorder=6)
    wx=-B*.75; wy=d-.35
    ax.annotate('',xy=(0,wy),xytext=(wx,wy),
                arrowprops=dict(arrowstyle='->',color='#f85149',lw=2.5),zorder=7)
    ax.text(wx-.5,wy+.3,f'P = {r["P"]:.1f} kN/m',fontsize=9,color='#f85149',ha='right',zorder=7)
    s_ok=r['FOS_s']>=1.25; o_ok=r['FOS_o']>=1.50; b_ok=r['q_max']<p['q_allow']
    summary=(f"Sliding FOS:      {r['FOS_s']:.3f}  {'✓' if s_ok else '✗'}\n"
             f"Overturning FOS:  {r['FOS_o']:.3f}  {'✓' if o_ok else '✗'}\n"
             f"q_max:            {r['q_max']:.0f} kPa  {'✓' if b_ok else '✗'}")
    ax.text(.02,.98,summary,transform=ax.transAxes,fontsize=9,va='top',ha='left',
            fontfamily='monospace',color='#e6edf3',
            bbox=dict(boxstyle='round,pad=.5',facecolor='#0d1117',edgecolor='#21262d',alpha=.95))
    handles=[mpatches.Patch(facecolor='#1e3a5f',alpha=.9,edgecolor='#2c5f9e',label='Caisson (Concrete)'),
             mpatches.Patch(facecolor='#C8A96E',alpha=.85,edgecolor='#8B6A38',label='Rubble Foundation'),
             mlines.Line2D([],[],color='#4493f8',ls='--',lw=2,label=f'Water Level (d = {d:.2f} m)')]
    leg=ax.legend(handles=handles,loc='upper right',fontsize=9,frameon=True,
                  facecolor='#0d1117',edgecolor='#1c2230',framealpha=.97)
    for t in leg.get_texts(): t.set_color('#d1d9e6')
    ax.set_xlabel('Distance  (m)',fontsize=11,color='#7d8590')
    ax.set_ylabel('Elevation  (m)',fontsize=11,color='#7d8590')
    ax.set_title('Vertical Caisson — Cross-Section & Stability',fontsize=13,fontweight='bold',color='#e6edf3',pad=12)
    ax.grid(True,alpha=.10,color='#2a3040',lw=.5)
    ax.tick_params(colors='#7d8590',labelsize=9)
    for sp in ax.spines.values(): sp.set_color('#1c2230')
    pad=B*1.3; ax.set_xlim(-pad,B+pad); ax.set_ylim(-rb_h-2.5,H_c+4.5)
    fig.tight_layout()
    return fig

# ════════════════════════════════════════════════════════════════════
#  KEYS
# ════════════════════════════════════════════════════════════════════
RM_KEYS=['rm_crest_elev','rm_seabed_elev','rm_crest_width','rm_upper_height',
         'rm_lower_height','rm_sw_upper','rm_sw_lower','rm_sw_berm',
         'rm_lw_upper','rm_lw_lower','rm_lw_berm','rm_dhw','rm_surcharge',
         'rm_num_slices','rm_n_layers',
         'rm_n1','rm_t1','rm_g1','rm_c1','rm_phi1',
         'rm_n2','rm_t2','rm_g2','rm_c2','rm_phi2',
         'rm_n3','rm_t3','rm_g3','rm_c3','rm_phi3',
         'rm_n4','rm_t4','rm_g4','rm_c4','rm_phi4',
         'rm_n5','rm_t5','rm_g5','rm_c5','rm_phi5',
         'rm_gr','rm_cr_','rm_phir']
CS_KEYS=['cs_B','cs_H_c','cs_gamma_c','cs_d','cs_H1pct','cs_gamma_w','cs_mu','cs_q_allow']
WU_KEYS=['wu_H','wu_T','wu_slope','wu_Rc','wu_gamma_f','wu_gamma_b','wu_beta']

def save_proj():
    d={k:st.session_state[k] for k in RM_KEYS+CS_KEYS+WU_KEYS if k in st.session_state}
    d['__v__']='4.0'; d['__ts__']=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return json.dumps(d,indent=2)
def load_proj(up):
    d=json.loads(up.read().decode())
    for k,v in d.items():
        if not k.startswith('__'): st.session_state[k]=v

# ════════════════════════════════════════════════════════════════════
#  UI HELPERS
# ════════════════════════════════════════════════════════════════════
def sec(label, icon=""):
    dot='<span class="sl-dot"></span>' if not icon else ''
    st.markdown(f'<div class="sec-lbl">{dot}{icon}&nbsp;{label}</div>',unsafe_allow_html=True)

def num(label,key,**kw):
    kw.setdefault('format','%.2f')
    v=st.number_input(label,value=float(st.session_state[key]),key=f'_n_{key}',**kw)
    st.session_state[key]=v; return v

def txt(label,key):
    st.session_state[key]=st.text_input(label,value=st.session_state[key],key=f'_t_{key}')
    return st.session_state[key]

def eq_block(step_label, formula, substitution, result, reference=None):
    ref_html=f'<div class="eq-ref">Reference: {reference}</div>' if reference else ''
    # Escape HTML special chars in substitution/formula
    sub_safe=substitution.replace('<','&lt;').replace('>','&gt;')
    formula_safe=formula.replace('<','&lt;').replace('>','&gt;')
    st.markdown(f"""<div class="eq-card">
  <div class="eq-step">{step_label}</div>
  <div class="eq-formula">{formula_safe}</div>
  <div class="eq-sub">{sub_safe}</div>
  <div class="eq-result">= {result}</div>
  {ref_html}
</div>""",unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
<div style="padding:16px 6px 10px 6px">
  <div style="font-size:24px;font-weight:900;color:#e6edf3;letter-spacing:-.5px">HYRCAN</div>
  <div style="font-size:11px;color:#7d8590;letter-spacing:1px;margin-top:1px">ENGINEERING SUITE  ·  v4.0</div>
</div>""",unsafe_allow_html=True)

    with st.expander("ℹ️  About",expanded=False):
        st.markdown("""
<div class="info-card" style="margin:0">
  <div class="ic-body">
    <b style="color:#e6edf3">Developer</b><br>
    Mathias Adjei Tawiah<br>
    <span style="font-size:11px">mathiasadjeitawiah@gmail.com</span><br><br>
    <b style="color:#e6edf3">Institution</b><br>
    Hohai University<br>
    <span style="font-size:11px">College of Harbor, Coastal &amp; Offshore Engineering</span><br><br>
    <b style="color:#e6edf3">Standards</b><br>
    <span class="badge">EurOtop 2018</span>
    <span class="badge">JTS 154-1-2011</span>
    <span class="badge">GB 50286-2013</span>
  </div>
</div>""",unsafe_allow_html=True)

    st.markdown('<div class="sec-lbl"><span class="sl-dot"></span>&nbsp;PROJECT FILE</div>',unsafe_allow_html=True)
    st.download_button('⬇  Export Project (.json)',data=save_proj(),
                       file_name='hyrcan_project.json',mime='application/json',
                       use_container_width=True)
    up=st.file_uploader('Import Project',type='json',label_visibility='collapsed')
    if up:
        load_proj(up); st.success('Project loaded.',icon='✅'); st.rerun()

    st.markdown("""
<div style="margin-top:20px;padding-top:14px;border-top:1px solid #1c2230" class="sb-link">
  <a href="https://github.com/Mathias-lab/hyrcan-suite" target="_blank">⬡  GitHub Repository</a><br>
  <a href="https://linkedin.com/in/mathias-adjei-tawiah" target="_blank">⬡  LinkedIn</a>
</div>
<div style="font-size:10px;color:#484f58;margin-top:14px">Coordinates verified against HYRCAN 3.0</div>
""",unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════
#  HEADER
# ════════════════════════════════════════════════════════════════════
st.markdown("""
<div style="border-bottom:1px solid #1c2230;padding:14px 0 13px 0;margin-bottom:6px;
            display:flex;align-items:center;gap:14px">
  <span style="font-size:28px">⚓</span>
  <div>
    <span style="font-size:20px;font-weight:800;color:#e6edf3;letter-spacing:-.3px">
      HYRCAN Engineering Suite</span>
    <span style="font-size:13px;color:#7d8590;margin-left:8px">v4.0</span>
    <div style="font-size:12px;color:#7d8590;margin-top:2px">
      Rubble Mound Stability &nbsp;·&nbsp; Vertical Caisson Analysis &nbsp;·&nbsp;
      Wave Run-Up &nbsp;·&nbsp; EurOtop 2018 &nbsp;·&nbsp; JTS 154-1-2011 &nbsp;·&nbsp; GB 50286-2013
    </div>
  </div>
</div>""",unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════
#  TABS
# ════════════════════════════════════════════════════════════════════
tab1,tab2,tab3,tab4=st.tabs([
    "🏗️  Rubble Mound",
    "🏛️  Vertical Caisson",
    "🌊  Wave Run-Up",
    "📋  Project History",
])

# ────────────────────────────────────────────────────────────────────
#  TAB 1
# ────────────────────────────────────────────────────────────────────
with tab1:
    left,right=st.columns([1,1.65],gap='large')
    with left:
        sec("EMBANKMENT GEOMETRY","📐")
        c1,c2=st.columns(2)
        with c1:
            ce=num('Crest elevation (m)','rm_crest_elev')
            cw=num('Crest width (m)','rm_crest_width',min_value=0.5)
            uh=num('Upper height (m)','rm_upper_height',min_value=0.0)
        with c2:
            se=num('Seabed elevation (m)','rm_seabed_elev')
            lh=num('Lower height (m)','rm_lower_height',min_value=0.0)
        total_h=ce-se
        if abs((uh+lh)-total_h)>0.01:
            st.error(f"Height mismatch: upper+lower = {uh+lh:.3f} m  ≠  crest−seabed = {total_h:.3f} m")

        sec("SLOPE RATIOS  (H : V)  —  berm = 0 for single slope","📊")
        c1,c2,c3=st.columns(3)
        with c1: sw_u=num('SW upper','rm_sw_upper',min_value=0.1); sw_l=num('SW lower','rm_sw_lower',min_value=0.1)
        with c2: lw_u=num('LW upper','rm_lw_upper',min_value=0.1); lw_l=num('LW lower','rm_lw_lower',min_value=0.1)
        with c3: sw_b=num('SW berm (m)','rm_sw_berm',min_value=0.0); lw_b=num('LW berm (m)','rm_lw_berm',min_value=0.0)
        st.caption(f"Seaward: **{'Single-slope' if sw_b==0 else f'Double (berm {sw_b:.1f} m)'}**"
                   f"  ·  Landward: **{'Single-slope' if lw_b==0 else f'Double (berm {lw_b:.1f} m)'}**")

        sec("HYDRAULIC CONDITIONS","🌊")
        c1,c2,c3=st.columns(3)
        with c1: dhw=num('DHW level (m)','rm_dhw')
        with c2: q=num('Surcharge (kN/m²)','rm_surcharge',min_value=0.0)
        with c3: ns=num('Slices','rm_num_slices',min_value=10.0,step=5.0,format='%.0f')

        sec("FOUNDATION LAYERS","🧱")
        n_layers=st.slider('Number of soil layers',min_value=3,max_value=5,
                            value=int(st.session_state['rm_n_layers']),step=1)
        st.session_state['rm_n_layers']=n_layers
        # Column headers
        hcols=st.columns([2.2,1.3,1.3,1.3,1.3])
        for col,h in zip(hcols,["Name","Thickness (m)","γ (kN/m³)","c (kPa)","φ (°)"]):
            col.markdown(f"<span style='font-size:11px;font-weight:600;color:#7d8590'>{h}</span>",unsafe_allow_html=True)
        lnames=[]; lthick=[]; lprops=[]
        for i in range(1,n_layers+1):
            cn,ct,cg,cc,cp=st.columns([2.2,1.3,1.3,1.3,1.3])
            with cn: nm=txt(' ',f'rm_n{i}')
            with ct: t=num(' ',f'rm_t{i}',min_value=0.1)
            with cg: gv=num(' ',f'rm_g{i}',min_value=10.0)
            with cc: cv=num(' ',f'rm_c{i}',min_value=0.0)
            with cp: phiv=num(' ',f'rm_phi{i}',min_value=0.0)
            lnames.append(nm); lthick.append(t); lprops.append((gv,cv,phiv))

        sec("RUBBLE MOUND PROPERTIES","⚙️")
        c1,c2,c3=st.columns(3)
        with c1: gr=num('γ (kN/m³)','rm_gr',min_value=10.0)
        with c2: cr_=num('c (kPa)','rm_cr_',min_value=0.0)
        with c3: phir=num('φ (°)','rm_phir',min_value=0.0)

        layers=compute_layer_elevations(se,lthick)
        sec("LAYER ELEVATION PREVIEW","📋")
        st.dataframe(pd.DataFrame({
            'Layer':lnames,
            'Top (m)':[f"{layers[f'top{i+1}']:.3f}" for i in range(n_layers)],
            'Bottom (m)':[f"{layers[f'bot{i+1}']:.3f}" for i in range(n_layers)],
            'Thickness (m)':[f"{layers[f'top{i+1}']-layers[f'bot{i+1}']:.3f}" for i in range(n_layers)],
            'HYRCAN Y-top':[f"{layers[f'Y_top{i+1}']:.3f}" for i in range(n_layers)],
            'HYRCAN Y-bot':[f"{layers[f'Y_bot{i+1}']:.3f}" for i in range(n_layers)],
        }),use_container_width=True,hide_index=True)
        st.caption(f"Datum shift: **+{layers['datum']:.3f} m**  ·  Bottom of '{lnames[-1]}' → Y = 0.000")
        st.markdown('<br>',unsafe_allow_html=True)
        gen=st.button('🚀  Generate HYRCAN Coordinates',type='primary',use_container_width=True)

    with right:
        layers=compute_layer_elevations(se,lthick)
        valid=abs((uh+lh)-(ce-se))<=0.01

        if gen and valid:
            g=compute_geometry(ce,se,cw,uh,lh,sw_u,sw_l,sw_b,lw_u,lw_l,lw_b,layers)
            st.session_state.update({'rm_g':g,'rm_layers':layers,'rm_lnames':lnames,
                'rm_lthick':lthick,'rm_lprops':lprops,'rm_nl':n_layers,
                'rm_sw_u':sw_u,'rm_sw_l':sw_l,'rm_sw_b':sw_b,
                'rm_lw_u':lw_u,'rm_lw_l':lw_l,'rm_lw_b':lw_b,'rm_generated':True})
            checks_=verify_geometry(g,ce,se,cw,uh,lh,sw_u,sw_l,sw_b,lw_u,lw_l,lw_b)
            push_history("Rubble Mound",f"Crest {ce} m / Seabed {se} m",
                {'crest_elev':ce,'seabed_elev':se,'crest_width':cw,'dhw':dhw,'surcharge':q},
                {'all_pass':all(c[3] for c in checks_),'model_width':round(g['model_right'],2),'datum':round(layers['datum'],3)})

        sec("CROSS-SECTION PREVIEW","🖼️")
        try:
            g_p=compute_geometry(ce,se,cw,uh,lh,sw_u,sw_l,sw_b,lw_u,lw_l,lw_b,layers)
            fp=draw_rubble(g_p,layers,dhw,ce,se,uh,lh,lnames,n_layers,q_val=q)
            st.pyplot(fp,use_container_width=True); plt.close(fp)
        except Exception as e:
            st.warning(f"Preview unavailable: {e}")

        if st.session_state.get('rm_generated'):
            g=st.session_state['rm_g']; layers_=st.session_state['rm_layers']
            lnames_=st.session_state['rm_lnames']; lprops_=st.session_state['rm_lprops']
            nl_=st.session_state['rm_nl']
            sw_u_=st.session_state.get('rm_sw_u',sw_u); sw_l_=st.session_state.get('rm_sw_l',sw_l)
            sw_b_=st.session_state.get('rm_sw_b',sw_b); lw_u_=st.session_state.get('rm_lw_u',lw_u)
            lw_l_=st.session_state.get('rm_lw_l',lw_l); lw_b_=st.session_state.get('rm_lw_b',lw_b)

            sec("GEOMETRY VERIFICATION","✅")
            checks=verify_geometry(g,ce,se,cw,uh,lh,sw_u_,sw_l_,sw_b_,lw_u_,lw_l_,lw_b_)
            all_ok=all(c[3] for c in checks)
            if all_ok: st.success('All geometry checks passed.',icon='✅')
            else: st.error('One or more checks failed — review parameters.')
            st.dataframe(pd.DataFrame({'Check':[c[0] for c in checks],
                'Expected':[f'{c[1]:.4f}' for c in checks],
                'Calculated':[f'{c[2]:.4f}' for c in checks],
                'Status':['Pass' if c[3] else 'FAIL' for c in checks]}),
                use_container_width=True,hide_index=True)
            pts_=g['pts']
            mc1,mc2,mc3,mc4=st.columns(4)
            mc1.metric('Model Width',f"{g['model_right']:.3f} m")
            mc2.metric('Crest Left X',f"{pts_['cl'][0]:.3f} m")
            mc3.metric('Crest Right X',f"{pts_['cr'][0]:.3f} m")
            mc4.metric('Datum Shift',f"+{layers_['datum']:.3f} m")

            # ── Step 7 Material Table (HTML — no ASCII) ────────────────
            sec("STEP 7 — MATERIAL PROPERTIES TABLE","📊")
            st.markdown(mat_table_html((gr,cr_,phir),lnames_,lprops_,nl_),unsafe_allow_html=True)
            st.markdown('<br>',unsafe_allow_html=True)

            sec("HYRCAN SETUP INSTRUCTIONS","🔧")
            instr=gen_instructions(g,layers_,dhw,q,ns,ce,se,uh,lh,
                sw_u_,sw_l_,sw_b_,lw_u_,lw_l_,lw_b_,cw,
                lnames_,lprops_,nl_,(gr,cr_,phir))
            st.markdown(f'<div class="code-out">{instr}</div>',unsafe_allow_html=True)

            sec("EXPORT","⬇️")
            # Build exports
            fig_buf=io.BytesIO()
            fe=draw_rubble(g,layers_,dhw,ce,se,uh,lh,lnames_,nl_,q_val=q)
            fe.savefig(fig_buf,format='png',dpi=300,bbox_inches='tight',facecolor='white')
            plt.close(fe); fig_bytes_=fig_buf.getvalue()
            coords_txt=gen_coords_txt(g,layers_,dhw,q,ce,se,nl_,sw_b_,lw_b_)
            e1,e2,e3,e4,e5=st.columns(5)
            with e1: st.download_button('📄 Coords (.txt)',data=coords_txt,file_name='hyrcan_coords.txt',mime='text/plain',use_container_width=True)
            with e2: st.download_button('📋 Instructions (.txt)',data=instr,file_name='hyrcan_instructions.txt',mime='text/plain',use_container_width=True)
            with e3:
                wb_=build_word_rubble(g,layers_,dhw,q,checks,ce,se,uh,lh,sw_u_,sw_l_,sw_b_,lw_u_,lw_l_,lw_b_,cw,lnames_,lprops_,nl_,(gr,cr_,phir),fig_bytes_) if HAS_DOCX else io.BytesIO()
                st.download_button('📝 Report (.docx)',data=wb_,file_name='hyrcan_rubble_report.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',use_container_width=True)
            with e4:
                xb_=build_excel_rubble(g,layers_,dhw,q,checks,ce,se,uh,lh,sw_u_,sw_l_,sw_b_,lw_u_,lw_l_,lw_b_,cw,lnames_,lprops_,nl_,(gr,cr_,phir)) if HAS_XL else io.BytesIO()
                st.download_button('📊 Data (.xlsx)',data=xb_,file_name='hyrcan_rubble_data.xlsx',mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',use_container_width=True)
            with e5: st.download_button('🖼️ Plot (PNG)',data=fig_bytes_,file_name='hyrcan_section.png',mime='image/png',use_container_width=True)
        elif gen and not valid:
            st.error('Correct the height mismatch before generating.')
        elif not st.session_state.get('rm_generated'):
            st.info('Enter parameters on the left and click Generate to produce HYRCAN output.',icon='ℹ️')

# ────────────────────────────────────────────────────────────────────
#  TAB 2
# ────────────────────────────────────────────────────────────────────
with tab2:
    l2,r2=st.columns([1,1.65],gap='large')
    with l2:
        sec("CAISSON GEOMETRY","📐")
        c1,c2=st.columns(2)
        with c1: B_cs=num('Width B (m)','cs_B',min_value=1.0); gc=num('γ_c (kN/m³)','cs_gamma_c',min_value=10.0)
        with c2: Hc=num('Height H_c (m)','cs_H_c',min_value=1.0)
        sec("HYDRAULIC CONDITIONS","🌊")
        c1,c2,c3=st.columns(3)
        with c1: d_=num('Depth d (m)','cs_d',min_value=0.1)
        with c2: H1=num('Wave H₁% (m)','cs_H1pct',min_value=0.1)
        with c3: gw=num('γ_w (kN/m³)','cs_gamma_w',min_value=9.0)
        sec("STABILITY PARAMETERS","⚙️")
        c1,c2=st.columns(2)
        with c1: mu=num('Friction μ','cs_mu',min_value=0.1,max_value=1.0,step=0.01)
        with c2: qa=num('q_allow (kPa)','cs_q_allow',min_value=50.0)
        st.markdown("""
<div class="info-card" style="margin-top:18px">
  <div class="ic-title">Design Standards</div>
  <div class="ic-body">
    Wave force: JTS 154-1-2011, Appendix A<br>
    Sliding FOS ≥ 1.25 — GB 50286-2013 §5.3.2<br>
    Overturning FOS ≥ 1.50 — GB 50286-2013 §5.3.3<br>
    Bearing: q_max &lt; q_allow — JTS 154-1-2011 §5.3.4
  </div>
</div>""",unsafe_allow_html=True)

    with r2:
        cs_p=dict(B=B_cs,H_c=Hc,gamma_c=gc,d=d_,H1pct=H1,gamma_w=gw,mu=mu,q_allow=qa)
        r=caisson_fos(**cs_p)
        s_ok=r['FOS_s']>=1.25; o_ok=r['FOS_o']>=1.50; b_ok=r['q_max']<qa

        sec("STABILITY SUMMARY","📊")
        m1,m2,m3=st.columns(3)
        m1.metric('Sliding FOS',f"{r['FOS_s']:.3f}",delta='≥ 1.25 Pass' if s_ok else '< 1.25 Fail',delta_color='normal' if s_ok else 'inverse')
        m2.metric('Overturning FOS',f"{r['FOS_o']:.3f}",delta='≥ 1.50 Pass' if o_ok else '< 1.50 Fail',delta_color='normal' if o_ok else 'inverse')
        m3.metric('Bearing q_max',f"{r['q_max']:.1f} kPa",delta=f"< {int(qa)} kPa Pass" if b_ok else f"≥ {int(qa)} kPa Fail",delta_color='normal' if b_ok else 'inverse')
        if s_ok and o_ok and b_ok: st.success('All stability criteria satisfied — design is adequate.',icon='✅')
        else: st.error('One or more criteria not satisfied — revise design parameters.',icon='❌')

        push_history("Vertical Caisson",f"B={B_cs} m / H_c={Hc} m",
            {'B':B_cs,'H_c':Hc,'gamma_c':gc,'d':d_,'H1pct':H1,'mu':mu},
            {'FOS_s':round(r['FOS_s'],4),'FOS_o':round(r['FOS_o'],4),
             'q_max':round(r['q_max'],2),'all_pass':bool(s_ok and o_ok and b_ok)})

        sec("RESULTS TABLE","📋")
        st.dataframe(pd.DataFrame({
            'Check':['Self-weight W','Wave Force P','Sliding FOS','Overturning FOS','Bearing q_max'],
            'Calculated':[f"{r['W']:,.1f} kN/m",f"{r['P']:,.3f} kN/m",
                          f"{r['FOS_s']:.3f}",f"{r['FOS_o']:.3f}",f"{r['q_max']:,.1f} kPa"],
            'Required':['—','—','≥ 1.25','≥ 1.50',f"< {int(qa)} kPa"],
            'Standard':['—','JTS 154-1-2011 App. A','GB 50286-2013 §5.3.2','GB 50286-2013 §5.3.3','JTS 154-1-2011 §5.3.4'],
            'Status':['—','—','Pass' if s_ok else 'FAIL','Pass' if o_ok else 'FAIL','Pass' if b_ok else 'FAIL']
        }),use_container_width=True,hide_index=True)

        sec("CALCULATION REPORT","🔢")
        eq_block("1 · Self-Weight",
            "W = B × H_c × γ_c",
            f"W = {B_cs:.2f} m  ×  {Hc:.2f} m  ×  {gc:.2f} kN/m³",
            f"{r['W']:,.2f}  kN/m")
        eq_block("2 · Wave Force",
            "P = ½ · γ_w · H₁%²  +  γ_w · d · H₁%",
            f"P = 0.5 × {gw:.2f} × {H1:.2f}²  +  {gw:.2f} × {d_:.2f} × {H1:.2f}\n"
            f"  = {0.5*gw*H1**2:.3f}  +  {gw*d_*H1:.3f}",
            f"{r['P']:,.3f}  kN/m",
            "JTS 154-1-2011, Appendix A")
        eq_block("3 · Sliding FOS",
            "FOS_s = μ · W / P",
            f"FOS_s = {mu:.3f} × {r['W']:,.2f}  /  {r['P']:,.3f}",
            f"{r['FOS_s']:.4f}   (required ≥ 1.25  →  {'Pass' if s_ok else 'FAIL'})",
            "GB 50286-2013 §5.3.2")
        eq_block("4 · Overturning FOS",
            "arm = d/2 + H₁%/3\nFOS_o = (W · B/2) / (P · arm)",
            f"arm = {d_:.3f}/2 + {H1:.3f}/3 = {r['arm']:.4f} m\n"
            f"M_res = {r['W']:,.2f} × {B_cs/2:.3f} = {r['M_res']:,.2f}  kN·m/m\n"
            f"M_ot  = {r['P']:,.3f} × {r['arm']:.4f} = {r['M_ot']:,.2f}  kN·m/m",
            f"{r['FOS_o']:.4f}   (required ≥ 1.50  →  {'Pass' if o_ok else 'FAIL'})",
            "GB 50286-2013 §5.3.3")
        eq_block("5 · Bearing Pressure",
            "q_max = W / B",
            f"q_max = {r['W']:,.2f}  /  {B_cs:.2f}",
            f"{r['q_max']:,.2f}  kPa   (allowable {int(qa)} kPa  →  {'Pass' if b_ok else 'FAIL'})",
            "JTS 154-1-2011 §5.3.4")

        sec("CROSS-SECTION","🖼️")
        fc=draw_caisson(cs_p,r)
        st.pyplot(fc,use_container_width=True); plt.close(fc)

        sec("EXPORT","⬇️")
        fig_c_buf=io.BytesIO()
        fc2=draw_caisson(cs_p,r)
        fc2.savefig(fig_c_buf,format='png',dpi=300,bbox_inches='tight',facecolor='white')
        plt.close(fc2); fc_bytes=fig_c_buf.getvalue()
        e1,e2,e3=st.columns(3)
        with e1:
            wc=build_word_caisson(cs_p,r,fc_bytes) if HAS_DOCX else io.BytesIO()
            st.download_button('📝 Report (.docx)',data=wc,file_name='hyrcan_caisson_report.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',use_container_width=True)
        with e2:
            xc=build_excel_caisson(cs_p,r) if HAS_XL else io.BytesIO()
            st.download_button('📊 Data (.xlsx)',data=xc,file_name='hyrcan_caisson_data.xlsx',mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',use_container_width=True)
        with e3: st.download_button('🖼️ Plot (PNG)',data=fc_bytes,file_name='hyrcan_caisson.png',mime='image/png',use_container_width=True)

# ────────────────────────────────────────────────────────────────────
#  TAB 3
# ────────────────────────────────────────────────────────────────────
with tab3:
    l3,r3=st.columns([1,1.5],gap='large')
    with l3:
        sec("WAVE PARAMETERS","🌊")
        c1,c2=st.columns(2)
        with c1: H_wu=num('H_m0 (m)','wu_H',min_value=0.1)
        with c2: T_wu=num('T_p (s)','wu_T',min_value=1.0)
        sec("STRUCTURE GEOMETRY","📐")
        slope_wu=num('Slope ratio H:V  (e.g. 2.5 → 1:2.5)','wu_slope',min_value=0.1)
        Rc_wu=num('Freeboard R_c (m)','wu_Rc',min_value=0.0)
        sec("REDUCTION FACTORS  — EurOtop 2018","⚙️")
        c1,c2=st.columns(2)
        with c1: gf_wu=num('Roughness γ_f','wu_gamma_f',min_value=0.1,max_value=1.0,step=0.01)
        with c2: gb_wu=num('Berm γ_b','wu_gamma_b',min_value=0.1,max_value=1.0,step=0.01)
        beta_wu=num('Obliquity β (°)','wu_beta',min_value=0.0,max_value=80.0,step=1.0,format='%.1f')
        with st.expander("Roughness Factor Reference  (EurOtop 2018, Table 5.2)"):
            st.dataframe(pd.DataFrame({'Structure Type':[
                'Smooth impermeable (concrete / asphalt)','Grass (short)','Grass (long/rough)',
                'Single-size rock (permeable)','Double-layer rock (permeable)',
                'Rock armour (rough permeable)','Tetrapods','Accropode','Xbloc',
                'Cube (single layer)','Cube (double layer)'],
                'γ_f':['1.00','1.00','0.90–1.00','0.55','0.45','0.40–0.55',
                       '0.38','0.46','0.45','0.47','0.50']}),
                use_container_width=True,hide_index=True)

    with r3:
        ta=1.0/slope_wu
        L0=calc_L0(T_wu); xi=calc_xi(ta,H_wu,L0)
        gbeta=calc_gbeta(beta_wu)
        Ru2=calc_Ru2(H_wu,xi,gf_wu,gb_wu,gbeta)
        Ru1=calc_Ru1(H_wu,xi,gf_wu,gb_wu,gbeta)
        Cr=calc_Cr(xi); q_wu=calc_q(H_wu,L0,xi,Rc_wu,gf_wu,gbeta)
        btype=breaker(xi); q_ls=q_wu*1000
        wu_in={'H':H_wu,'T':T_wu,'slope':slope_wu,'Rc':Rc_wu,'gf':gf_wu,'gb':gb_wu,'beta':beta_wu}
        wu_res={'L0':L0,'xi':xi,'gbeta':gbeta,'Ru2':Ru2,'Ru1':Ru1,'Cr':Cr,'q':q_wu,'btype':btype,'Rc':Rc_wu}
        push_history("Wave Run-Up",f"H={H_wu} m / T={T_wu} s / Rc={Rc_wu} m",
            wu_in,{'Ru2':round(Ru2,4),'Ru1':round(Ru1,4),'xi':round(xi,4),'q_L_s_m':round(q_ls,5),'breaker':btype})

        sec("RESULTS","📊")
        m1,m2,m3=st.columns(3)
        m1.metric("L₀  (m)",f"{L0:.3f}"); m2.metric("Iribarren ξ₀",f"{xi:.4f}"); m3.metric("Breaker",btype)
        m4,m5,m6=st.columns(3)
        m4.metric("R_u2%  (m)",f"{Ru2:.4f}"); m5.metric("R_u1%  (m)",f"{Ru1:.4f}"); m6.metric("C_r",f"{Cr:.4f}")
        m7,m8,m9=st.columns(3)
        m7.metric("γ_β",f"{gbeta:.4f}")
        m8.metric("R_c / R_u2%",f"{Rc_wu/Ru2:.3f}" if Ru2>0 else "—")
        m9.metric("q  (L/s/m)",f"{q_ls:.5f}" if q_ls<1 else f"{q_ls:.3f}")
        if q_ls<0.01: st.success(f"Overtopping = {q_ls:.5f} L/s/m — Negligible (< 0.01 L/s/m).",icon='✅')
        elif q_ls<1.0: st.success(f"Overtopping = {q_ls:.5f} L/s/m — Within acceptable limits.",icon='✅')
        elif q_ls<10.0: st.warning(f"Overtopping = {q_ls:.4f} L/s/m — Verify tolerable limits.")
        else: st.error(f"Overtopping = {q_ls:.3f} L/s/m — Excessive. Increase freeboard or armour roughness.")

        sec("CALCULATION REPORT  — EurOtop 2018","🔢")
        Ru_raw=1.65*gb_wu*gf_wu*gbeta*xi*H_wu
        cap_val=(4.0-1.5/max(math.sqrt(gf_wu*gbeta*xi),1e-9))*gb_wu*H_wu

        eq_block("Step 1 · Deep-Water Wavelength",
            "L₀ = g · T_p² / (2π)",
            f"L₀ = 9.81 × {T_wu:.2f}² / (2π)",
            f"{L0:.4f}  m")
        eq_block("Step 2 · Iribarren Number  (Surf-Similarity)",
            "ξ₀ = tan(α) / √(H_m0 / L₀)",
            f"ξ₀ = {ta:.5f} / √({H_wu:.3f} / {L0:.3f})",
            f"{xi:.4f}   →   Breaker class: {btype}")
        eq_block("Step 3 · Obliquity Reduction Factor",
            "γ_β = max(1 − 0.0033 · |β|,  0.736)",
            f"γ_β = max(1 − 0.0033 × {beta_wu:.1f},  0.736)",
            f"{gbeta:.5f}",
            "EurOtop 2018, §5.1.3")
        eq_block("Step 4 · Wave Run-Up R_u2%",
            "R_u2% = min( 1.65·γ_b·γ_f·γ_β·ξ₀·H_m0,  (4.0−1.5/√(γ_f·γ_β·ξ₀))·γ_b·H_m0 )",
            f"Formula = 1.65 × {gb_wu:.3f} × {gf_wu:.3f} × {gbeta:.4f} × {xi:.4f} × {H_wu:.3f} = {Ru_raw:.4f} m\n"
            f"Cap     = (4.0 − 1.5/√({gf_wu:.3f}×{gbeta:.4f}×{xi:.4f})) × {gb_wu:.3f} × {H_wu:.3f} = {cap_val:.4f} m",
            f"R_u2% = {Ru2:.4f} m   ·   R_u1% = 1.4 × R_u2% = {Ru1:.4f} m",
            "EurOtop 2018, Equation 5.2")
        eq_block("Step 5 · Reflection Coefficient",
            "C_r = 0.1 · ξ₀²   (capped at 1.0)",
            f"C_r = 0.1 × {xi:.4f}² = {0.1*xi**2:.5f}",
            f"{Cr:.5f}",
            "Postma (1989)")
        eq_block("Step 6 · Mean Overtopping Discharge",
            "q = √(g·H³)/√(L₀/H) · 0.2 · exp(−2.3·R_c/(γ_f·γ_β·H·ξ₀))",
            f"q = {q_wu:.7f}  m³/s/m",
            f"{q_ls:.6f}  L/s/m",
            "EurOtop 2018, Equation 5.6")

        sec("EXPORT","⬇️")
        e1,e2=st.columns(2)
        with e1:
            ww=build_word_runup(wu_in,wu_res) if HAS_DOCX else io.BytesIO()
            st.download_button('📝 Report (.docx)',data=ww,file_name='hyrcan_runup_report.docx',mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',use_container_width=True)
        with e2:
            xw=build_excel_runup(wu_in,wu_res) if HAS_XL else io.BytesIO()
            st.download_button('📊 Data (.xlsx)',data=xw,file_name='hyrcan_runup_data.xlsx',mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',use_container_width=True)

# ────────────────────────────────────────────────────────────────────
#  TAB 4 — PROJECT HISTORY
# ────────────────────────────────────────────────────────────────────
with tab4:
    history=load_history()
    hc,bc=st.columns([5,1])
    with hc:
        st.markdown(f"""<div style="font-size:13px;color:#7d8590;margin-bottom:12px">
  Session log — <b style="color:#e6edf3">{len(history)}</b> records
  (capacity: 60 entries, FIFO). Records persist for the duration of the server session.
</div>""",unsafe_allow_html=True)
    with bc:
        if st.button('🗑  Clear',use_container_width=True):
            save_history([]); st.rerun()

    if not history:
        st.info('No records yet. Run any calculation to begin logging.',icon='ℹ️')
    else:
        modules=["All"]+sorted({e['module'] for e in history})
        sel=st.selectbox('Filter by module',modules)
        filtered=[e for e in history if sel=="All" or e['module']==sel]

        for entry in filtered:
            res=entry.get('results',{})
            # Status indicator
            if 'all_pass' in res or 'all_checks_passed' in res:
                ok=res.get('all_pass',res.get('all_checks_passed',False))
                badge=f'<span class="p-pass">PASS</span>' if ok else '<span class="p-fail">FAIL</span>'
            elif 'Ru2' in res:
                badge=f'<span class="p-warn">{res.get("breaker","—")}</span>'
            else: badge=""

            with st.expander(f"**[{entry['module']}]**  {entry['label']}  ·  {entry['ts']}  {badge}"):
                c1,c2=st.columns(2)
                with c1:
                    st.markdown("**Input Parameters**")
                    st.dataframe(pd.DataFrame([[k,str(v)] for k,v in entry.get('params',{}).items()],
                        columns=['Parameter','Value']),use_container_width=True,hide_index=True)
                with c2:
                    st.markdown("**Results**")
                    st.dataframe(pd.DataFrame([[k,str(v)] for k,v in res.items()],
                        columns=['Result','Value']),use_container_width=True,hide_index=True)
                st.download_button(f'Export Record  [{entry["id"]}]  (.json)',
                    data=json.dumps(entry,indent=2),
                    file_name=f'hyrcan_record_{entry["id"]}.json',
                    mime='application/json',key=f'dl_{entry["id"]}')

        st.markdown('<br>',unsafe_allow_html=True)
        sec("BATCH EXPORT","📦")
        if HAS_XL:
            wb_h=openpyxl.Workbook(); ws_h=wb_h.active; ws_h.title="History"
            ws_h.sheet_view.showGridLines=False
            title_f_=Font(name='Calibri',bold=True,size=14,color='003366')
            ws_h['A1']="HYRCAN v4.0 — Calculation History Log"; ws_h['A1'].font=title_f_
            ws_h.merge_cells('A1:E1')
            _xl_write_table(ws_h,["Timestamp","Module","Label","Parameters","Results"],
                [[e['ts'],e['module'],e['label'],json.dumps(e.get('params',{})),json.dumps(e.get('results',{}))]
                 for e in history],sr=3)
            ws_h.column_dimensions['A'].width=22; ws_h.column_dimensions['B'].width=20
            ws_h.column_dimensions['C'].width=32; ws_h.column_dimensions['D'].width=50
            ws_h.column_dimensions['E'].width=50
            hxbuf=io.BytesIO(); wb_h.save(hxbuf); hxbuf.seek(0)
        hb1,hb2=st.columns(2)
        with hb1:
            st.download_button('Export All Records (.json)',data=json.dumps(history,indent=2),
                file_name='hyrcan_history.json',mime='application/json',use_container_width=True)
        with hb2:
            if HAS_XL:
                st.download_button('Export All Records (.xlsx)',data=hxbuf,
                    file_name='hyrcan_history.xlsx',
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    use_container_width=True)

# ════════════════════════════════════════════════════════════════════
#  FOOTER
# ════════════════════════════════════════════════════════════════════
st.markdown("""
<div style="border-top:1px solid #1c2230;margin-top:32px;padding-top:14px;
            text-align:center;font-size:11px;color:#484f58;line-height:2.2">
  HYRCAN Engineering Suite v4.0 &nbsp;·&nbsp;
  Mathias Adjei Tawiah &nbsp;·&nbsp;
  mathiasadjeitawiah@gmail.com &nbsp;·&nbsp;
  Hohai University — College of Harbor, Coastal and Offshore Engineering<br>
  Standards: EurOtop 2018 &nbsp;·&nbsp; JTS 154-1-2011 &nbsp;·&nbsp; GB 50286-2013
  &nbsp;·&nbsp; Coordinate logic verified against HYRCAN 3.0
</div>
""",unsafe_allow_html=True)
